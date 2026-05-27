from __future__ import annotations

import re
import json
import tempfile
import unittest
import copy
import xml.etree.ElementTree as ET
from datetime import date, datetime
from pathlib import Path
from collections import Counter
from zipfile import ZipFile

from docx import Document

ROOT = Path(__file__).resolve().parents[1]

import sys

if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_payload import (
    build_payload,
    derive_preface,
    derive_project_background,
    normalize_survey_period,
    parse_markdown_content,
    validate_payload,
)
from scripts.cluster_dimensions import EFFICACY_GROUPING, cluster_dimensions
from scripts.expression_data import (
    MAX_ANALYSIS_CHARS,
    MIN_ANALYSIS_CHARS,
    OPENING_STYLE_LIBRARY,
    build_analysis_paragraph,
    is_complete_analysis,
)
from scripts.final_validate_docx import FinalValidationError, validate_docx
from scripts.parse_questionnaire import parse_sheet
from scripts.render_from_template import TemplateRenderer


class Namespace:
    def __init__(self, **kwargs):
        self.__dict__.update(kwargs)


def analysis_opening(text: str) -> str:
    match = re.match(r"^([^，。；;,.]{2,32})[，。；;,.]", str(text or ""))
    if match:
        return match.group(1)
    return str(text or "")[:16]


def expected_result_table_questions(payload: dict) -> list[str]:
    qmap = {
        question["question_ref"]: question["question"]
        for question in payload["attachments"]["attachment1_questions"]
    }
    expected: list[str] = []
    seen: set[str] = set()
    for section in payload["result_analysis"]["sections"]:
        for subtopic in section["subtopics"]:
            for question_ref in subtopic["question_refs"]:
                if question_ref in qmap and question_ref not in seen:
                    expected.append(qmap[question_ref])
                    seen.add(question_ref)
    return expected


AI_KEY_ISSUE_PARAGRAPH_1 = (
    "血压控制相关重点问题反映出患者对厄贝沙坦氢氯噻嗪片真实使用效果的核心判断。"
    "从样本反馈看，正向评价集中度较高，说明多数患者能够在日常监测和自身感受中形成较稳定的控压认知，"
    "产品使用体验具备较好的患者基础。与此同时，少数患者仍然表现出对血压波动的谨慎感受，"
    "这类反馈提示后续沟通不能只停留在总体疗效呈现，还需要结合复诊记录、家庭监测和用药执行情况进行解释。"
    "对于区域样本而言，这一问题还能够帮助判断患者是否真正理解治疗目标，以及是否能够把血压变化与规范服药联系起来。"
    "综合判断，该问题的管理价值在于把已有正向体验转化为更稳定的长期信任，并识别需要持续随访的波动样本。"
)


AI_KEY_ISSUE_PARAGRAPH_2 = (
    "不良反应相关重点问题主要体现患者对长期用药耐受性的实际感知。"
    "当前反馈显示，多数患者并未将头晕头痛等不适视为持续用药的主要障碍，说明安全性体验整体较为平稳，"
    "也为后续维持治疗连续性提供了基础。需要关注的是，低频不适虽然没有形成普遍压力，"
    "但在个体患者中仍可能影响服药信心和复诊沟通质量，尤其当患者缺少明确解释时，轻微症状也可能被放大为停药顾虑。"
    "因此，报告需要把主流稳定体验和少数敏感反馈同时呈现，避免用平均判断掩盖个体管理需求。"
    "因此，该重点问题应被用于识别敏感人群和优化说明方式，使患者能够在理解风险边界的基础上保持规范用药。"
)


AI_ADHERENCE_KEY_ISSUE_1 = (
    "漏服应对问题集中反映了患者在真实用药场景中的执行稳定性。"
    "从反馈结构看，多数患者能够选择较为合理的补救方式，说明基础依从意识已经形成，"
    "但仍有部分患者存在跳过剂量、加倍补服或处理方式不确定等表现，这类行为虽然不是主流，"
    "却可能在长期治疗中累积为疗程中断和剂量波动风险。该问题的重点不在于简单评价患者是否听从医嘱，"
    "而在于识别患者面对突发漏服时是否具备清晰、可执行的处理规则。综合判断，后续分析应围绕标准化漏服指引和提醒支持展开，"
    "同时关注不同年龄和信息来源患者的理解差异，帮助患者把正确认知落实到具体行动之中。"
)


AI_ADHERENCE_KEY_ISSUE_2 = (
    "依从性提升需求体现了患者对长期用药支持体系的真实期待。"
    "样本反馈显示，健康教育资料、副作用解释和提醒工具等需求较为突出，说明患者并不只是被动接受治疗安排，"
    "而是希望获得更清晰、更连续的外部支持。对于心达康胶囊这类需要稳定服用的品种而言，"
    "依从性问题往往同时受到药物认知、用药体验和日常提醒方式影响，单一宣教很难完全解决。"
    "如果支持内容缺少场景化解释，患者即使知道需要规律服药，也可能在症状变化或生活节奏变化时出现执行松动。"
    "综合判断，该重点问题应作为后续患者管理优化的入口，通过分层教育、风险解释和提醒机制协同提升长期用药连续性。"
)


def efficacy_questionnaire() -> dict:
    questions = [
        "您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？",
        "您服用厄贝沙坦氢氯噻嗪片期间，是否出现过头晕、头痛症状？",
        "您服用厄贝沙坦氢氯噻嗪片期间，是否出现过口渴、多尿情况？",
        "您通常通过什么渠道购买厄贝沙坦氢氯噻嗪片？",
        "您认为厄贝沙坦氢氯噻嗪片的价格对您的用药负担影响如何？",
        "您服用厄贝沙坦氢氯噻嗪片的剂量是否严格按照医嘱？",
        "您对厄贝沙坦氢氯噻嗪片的服药频率满意度如何？",
        "您服用厄贝沙坦氢氯噻嗪片期间，是否同时服用其他降压药物？",
        "您是否咨询过医生，确认厄贝沙坦氢氯噻嗪片与其他药物可同时服用？",
        "您所在地区的厄贝沙坦氢氯噻嗪片供应是否稳定？",
        "您服用厄贝沙坦氢氯噻嗪片期间，是否定期监测血压？",
        "您对厄贝沙坦氢氯噻嗪片药品包装的便利性评价如何？",
        "您对厄贝沙坦氢氯噻嗪片说明书中 “用法用量”“不良反应” 的标注清晰度评价如何？",
        "您认为当前获取的用药指导是否足够详细准确？",
    ]
    options = [
        ("A", "选项A", "700", "39.13%"),
        ("B", "选项B", "600", "33.54%"),
        ("C", "选项C", "300", "16.77%"),
        ("D", "选项D", "189", "10.56%"),
    ]
    data = []
    for i, q in enumerate(questions, start=1):
        data.append(
            {
                "number": i,
                "question": q,
                "total": "1789",
                "options": [{"label": a, "text": b, "count": c, "pct": d} for a, b, c, d in options],
            }
        )
    return {"question_count": len(data), "questions": data}


def adherence_questionnaire() -> dict:
    questions = [
        "您对心达康胶囊作用的了解程度如何？",
        "您获取心达康胶囊相关信息的主要来源是什么？",
        "您忘记服药后通常会怎么做？",
        "您是否会记录每日服药情况？",
        "您通常通过什么方式提醒自己服药？",
        "您是否能够坚持规律服用心达康胶囊？",
        "您是否愿意参加心达康胶囊相关健康教育活动？",
        "您希望获得哪些形式的健康教育支持？",
    ]
    options = [
        ("A", "非常了解", "410", "41.00%"),
        ("B", "比较了解", "360", "36.00%"),
        ("C", "一般了解", "180", "18.00%"),
        ("D", "不了解", "50", "5.00%"),
    ]
    data = []
    for i, q in enumerate(questions, start=1):
        data.append(
            {
                "number": i,
                "question": q,
                "total": "1000",
                "options": [{"label": a, "text": b, "count": c, "pct": d} for a, b, c, d in options],
            }
        )
    return {"question_count": len(data), "questions": data}


def sample_markdown() -> str:
    return f"""---
product: 厄贝沙坦氢氯噻嗪片
region: 北京市
time: 2025.10
attachment_name: 厄贝沙坦氢氯噻嗪片用药体验与疗效反馈患者调查问卷
survey_period: 2025年10月01日——2025年10月31日
valid_count: 1789
disclaimer_unit: 北京玖麟空科技有限公司
---

## 前言

高血压作为我国常见慢性非传染性疾病之一，其长期规范管理对于降低心脑血管并发症风险、改善患者生活质量具有重要意义。厄贝沙坦氢氯噻嗪片作为临床常用复方降压药物，在稳定控压、简化治疗方案及提升长期管理可执行性方面具有较明确的应用价值。随着患者对疗效感知、安全性体验和持续用药支持要求不断提高，仅依赖传统临床观察已难以全面反映真实使用场景中的需求差异。

本次调研聚焦北京市使用厄贝沙坦氢氯噻嗪片的患者群体，共收集有效问卷1789份，围绕药品疗效、安全性、用药行为、便利性、经济性、可及性及信息支持等核心维度展开系统分析，旨在从患者视角识别真实用药过程中的优势表现与待优化环节，同时关注区域化支持需求差异。通过对区域样本反馈的系统梳理，本报告将为企业后续优化患者教育内容、完善用药支持服务和提升产品应用沟通质量提供数据参考，也为临床端进一步加强慢病管理中的用药指导提供依据。

## 项目背景

随着人口老龄化加快和生活方式变化，高血压患病率持续维持在较高水平，患者对长期规范用药的依赖程度不断增强。厄贝沙坦氢氯噻嗪片作为复方降压药物，在真实临床应用中不仅承担控压作用，也直接影响患者对治疗便利性、耐受性与持续管理信心的判断。北京市患者就医渠道丰富、健康管理意识较强，为观察真实用药体验提供了较好的样本基础。

但现阶段围绕该药物在区域患者中的实际使用反馈、依从行为差异及服务支持需求，仍缺乏持续、细化的专项调研结果，导致产品优化与患者沟通改进缺少足够的本地化依据。本次调研因此聚焦患者真实用药场景，希望通过结构化问卷补足关键数据缺口，为后续产品服务优化和患者管理支持提供更具针对性的参考。

## 项目开展情况

本项目采用患者问卷方式开展。

## 问卷说明

本次分析采用结构化处理方式。

## 问卷结果分析

本次问卷从7个维度展开统计分析。

### 4.1药品疗效

本维度用于观察患者对控压效果的主观感知。

#### 血压控制效果分析

多数患者在该题中的反馈集中在正向选项，说明控压结果的主流感知较为稳定。

### 4.2药品安全性

本维度用于观察患者对常见不适反应的主观感知。

#### 头晕头痛不良反应分析

多数患者并未将该症状视为主要负担。

#### 口渴多尿不良反应分析

患者对该类不适的整体反馈同样偏轻。

### 4.3用药行为与习惯

本维度用于观察患者的长期行为执行情况。

#### 购药渠道分析

渠道选择主要集中于正规渠道。

#### 剂量遵循情况分析

多数患者能够保持较稳定的剂量执行。

#### 联合用药情况分析

部分患者存在联合用药需求。

#### 血压监测频率分析

规律监测行为仍有继续提升空间。

### 4.4用药便利性

本维度用于观察频率与包装体验。

#### 服药频率满意度分析

大多数患者认为服药频率可以接受。

#### 药品包装便利性分析

包装便利性整体较好。

### 4.5药品经济性

本维度用于观察价格感知。

#### 价格负担影响分析

多数患者并未将价格视为主要障碍。

### 4.6药品可及性

本维度用于观察供应稳定性。

#### 药品供应稳定性分析

大多数患者认为购药较为顺畅。

### 4.7用药指导信息评价

本维度用于观察说明信息和指导支持。

#### 说明书信息清晰度分析

大多数患者能够理解说明书中的核心信息。

#### 用药指导详细准确性分析

现有指导信息总体能够满足基础使用需求。

## 调研结果

### 5.1问卷重点问题分析

{AI_KEY_ISSUE_PARAGRAPH_1}

{AI_KEY_ISSUE_PARAGRAPH_2}

### 5.2调研结果分析

整体来看，患者对疗效、安全性和便利性的反馈基础较好。

### 5.3建议

1. 补充场景化患者教育材料。
"""


def minimal_markdown(product: str = "心达康胶囊", region: str = "安徽省", include_key_issue: bool = True) -> str:
    key_issue = (
        f"""
## 调研结果

### 5.1问卷重点问题分析

{AI_ADHERENCE_KEY_ISSUE_1}

{AI_ADHERENCE_KEY_ISSUE_2}
"""
        if include_key_issue
        else ""
    )
    return f"""---
product: {product}
region: {region}
survey_period: 2026年5月13日至5月31日
valid_count: 1000
disclaimer_unit: 北京玖麟空科技有限公司
---
{key_issue}
"""


def dynamic_ai_dimensions() -> dict:
    return {
        "dimensions": [
            {
                "name": "药物认知与信息获取",
                "intro": "本维度用于观察患者对药物作用的了解程度及信息来源结构。",
                "subtopics": [
                    {"patterns": ["作用", "了解", "认知"], "subtitle": "药物认知情况分析"},
                    {"patterns": ["获取", "来源", "说明书"], "subtitle": "信息获取来源分析"},
                ],
                "charts": [
                    {"patterns": ["作用"], "chart_type": "pie", "chart_style_profile": "efficacy_pie"},
                ],
            },
            {
                "name": "用药行为与依从性",
                "intro": "本维度用于观察患者用药执行和坚持情况。",
                "subtopics": [
                    {"patterns": ["剂量", "服药", "忘记|漏服"], "subtitle": "剂量执行与漏服应对分析"},
                    {"patterns": ["监测", "记录"], "subtitle": "自我管理行为分析"},
                    {"patterns": ["提醒", "依从|坚持规律"], "subtitle": "依从与提醒支持分析"},
                ],
                "charts": [
                    {"patterns": ["剂量"], "chart_type": "bar3d", "chart_style_profile": "behavior_bar"},
                ],
            },
            {
                "name": "健康教育与支持需求",
                "intro": "本维度用于观察患者对健康教育的接受意愿。",
                "subtopics": [
                    {"patterns": ["讲座", "培训", "参加", "兴趣", "希望|获得"], "subtitle": "健康教育参与意愿分析"},
                ],
                "charts": [],
            },
        ]
    }


def copy_docx_with_document_xml_replace(source: Path, target: Path, old: str, new: str, count: int = 1) -> None:
    with ZipFile(source) as src, ZipFile(target, "w") as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/document.xml":
                data = data.decode("utf-8", "ignore").replace(old, new, count).encode("utf-8")
            dst.writestr(item, data)


def chart_refs_between(docx_path: Path, start_text: str, end_text: str) -> list[str]:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    with ZipFile(docx_path) as zipped:
        root = ET.fromstring(zipped.read("word/document.xml"))
        rels = ET.fromstring(zipped.read("word/_rels/document.xml.rels"))
    relmap = {rel.attrib["Id"]: rel.attrib.get("Target", "") for rel in rels}
    inside = False
    refs: list[str] = []
    for paragraph in root.findall(".//w:body/w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
        if text == start_text:
            inside = True
            continue
        if text == end_text:
            break
        if inside:
            for chart in paragraph.findall(".//c:chart", ns):
                refs.append(relmap.get(chart.attrib.get(f"{{{ns['r']}}}id"), ""))
    return refs


def chart_part_type(docx_path: Path, chart_target: str) -> str:
    chart_name = f"word/{chart_target}" if not chart_target.startswith("word/") else chart_target
    ns = {"c": "http://schemas.openxmlformats.org/drawingml/2006/chart"}
    with ZipFile(docx_path) as zipped:
        root = ET.fromstring(zipped.read(chart_name))
    chart_types = []
    for elem in root.iter():
        tag = elem.tag.rsplit("}", 1)[-1]
        if tag.endswith("Chart") and tag not in {"chart", "plotArea"}:
            chart_types.append(tag)
    return chart_types[0] if chart_types else ""


class PipelineTest(unittest.TestCase):
    def test_cluster_dimensions_uses_efficacy_template(self) -> None:
        grouped = cluster_dimensions(efficacy_questionnaire())
        self.assertEqual(grouped["template_type"], "用药体验与疗效反馈")
        self.assertEqual(grouped["dimension_count"], 7)
        self.assertEqual(len(grouped["project_dimensions"]), 7)
        self.assertIn("用药指导信息评价", grouped["project_dimensions"])
        self.assertEqual(grouped["sections"][0]["section_title"], "药品疗效")
        self.assertEqual(grouped["sections"][0]["section_intro"], "本维度用于观察患者对控压效果的主观感知。")
        self.assertGreater(len(grouped["sections"]), 0)
        section_1 = grouped["sections"][0]
        self.assertGreater(len(section_1["subtopics"]), 0)
        subtopic_1 = section_1["subtopics"][0]
        self.assertIn("subtopic_index", subtopic_1)
        self.assertIn("question_refs", subtopic_1)
        self.assertEqual(subtopic_1["subtitle"], "血压控制效果分析")
        self.assertGreater(len(subtopic_1["question_refs"]), 0)

    def test_build_payload_matches_template_slots(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )
        validate_payload(payload)
        self.assertEqual(len(payload["preface"]), 2)
        self.assertEqual(len(payload["project_background"]), 2)
        for paragraph in payload["preface"]:
            self.assertNotRegex(paragraph, r"#+")
        for paragraph in payload["project_background"]:
            self.assertNotRegex(paragraph, r"#+")
        self.assertGreaterEqual(sum(len(p) for p in payload["preface"]), 380)
        self.assertLessEqual(sum(len(p) for p in payload["preface"]), 430)
        self.assertGreaterEqual(sum(len(p) for p in payload["project_background"]), 180)
        self.assertLessEqual(sum(len(p) for p in payload["project_background"]), 360)
        self.assertIn("北京市", "".join(payload["preface"]))
        self.assertIn("患者", "".join(payload["preface"]))
        self.assertIn("提供", payload["preface"][-1])
        self.assertIn("北京市", "".join(payload["project_background"]))
        self.assertIn("本次调研", "".join(payload["project_background"]))
        self.assertEqual(
            payload["project_execution"]["lines"],
            [
                "项目工具：调查问卷，14道选择题。其内容涵盖药品疗效、药品安全性、用药行为与习惯、用药便利性、药品经济性、药品可及性、用药指导信息评价7大维度，全面覆盖患者用药全流程关键节点。",
                "样本采集范围：北京市",
                "样本采集数量：本次共收集筛选到有效问卷1789份。",
                "样本采集时间：2025年10月01日——2025年10月31日",
            ],
        )
        self.assertEqual(payload["questionnaire_note"]["intro"], "为提升数据可比性与分析结果的科学性，本研究采用以下标准化统计处理方法对原始问卷数据进行系统规整：")
        self.assertEqual(len(payload["questionnaire_note"]["items"]), 4)
        self.assertEqual(
            payload["result_analysis"]["intro"],
            ["本次调查采用问卷调查方式，共有14个问题，本报告从药品疗效、药品安全性、用药行为与习惯、用药便利性、药品经济性、药品可及性、用药指导信息评价等7个维度展开统计分析。"],
        )
        self.assertEqual(len(payload["result_analysis"]["overview_charts"]), 2)
        self.assertEqual(payload["result_analysis"]["overview_charts"][0]["chart_type"], "pie")
        self.assertEqual(payload["result_analysis"]["overview_charts"][1]["chart_type"], "bar")
        self.assertEqual(payload["result_analysis"]["overview_charts"][0]["render_mode"], "image")
        self.assertEqual(payload["result_analysis"]["overview_charts"][1]["render_mode"], "image")
        self.assertEqual(
            payload["result_analysis"]["overview_charts"][0]["categories"],
            ["药品疗效", "药品安全性", "用药行为与习惯", "用药便利性", "药品经济性", "药品可及性", "用药指导信息评价"],
        )
        self.assertEqual(payload["result_analysis"]["overview_charts"][0]["categories"], payload["result_analysis"]["overview_charts"][1]["categories"])
        self.assertEqual(payload["result_analysis"]["overview_charts"][0]["values"], [1, 2, 5, 2, 1, 1, 2])
        self.assertEqual(payload["result_analysis"]["overview_charts"][0]["values"], payload["result_analysis"]["overview_charts"][1]["values"])
        self.assertEqual(payload["result_analysis"]["sections"][0]["section_title"], "药品疗效")
        self.assertEqual(payload["result_analysis"]["sections"][0]["section_intro"], ["本维度用于观察患者对控压效果的主观感知。"])
        self.assertEqual(payload["result_analysis"]["sections"][0]["subtopics"][0]["subtitle"], "血压控制效果分析")
        self.assertEqual(payload["result_analysis"]["sections"][0]["visual_groups"][0]["chart_type"], "pie")
        self.assertEqual(len(payload["result_analysis"]["sections"][3]["visual_groups"]), 2)
        self.assertEqual(len(payload["result_analysis"]["sections"][4]["visual_groups"]), 1)
        self.assertEqual(len(payload["result_analysis"]["sections"][5]["visual_groups"]), 1)
        self.assertEqual(len(payload["result_analysis"]["sections"][6]["visual_groups"]), 2)
        self.assertIsNone(payload["result_analysis"]["sections"][3]["visual_groups"][0]["chart_ref"])
        self.assertEqual(
            payload["result_analysis"]["sections"][3]["visual_groups"][0]["table_data"]["question"],
            "您对厄贝沙坦氢氯噻嗪片的服药频率满意度如何？",
        )
        self.assertEqual(payload["report_title"], "问卷调研分析报告")
        self.assertEqual(payload["header_text"], "问卷调研分析报告")
        self.assertEqual(payload["meta"]["survey_period"], "2025年10月01日——2025年10月31日")
        self.assertEqual(payload["service"]["unit"], "北京玖麟空科技有限公司")
        self.assertRegex(payload["service"]["date"], r"^2025年11月\d{2}日$")
        service_dt = datetime.strptime(payload["service"]["date"], "%Y年%m月%d日").date()
        self.assertGreaterEqual(service_dt, date(2025, 11, 1))
        self.assertLessEqual(service_dt, date(2025, 11, 30))
        self.assertEqual(payload["disclaimer"]["unit"], payload["service"]["unit"])
        self.assertEqual(payload["disclaimer"]["date"], payload["service"]["date"])
        self.assertEqual(payload["summary"]["key_issue_analysis"], [AI_KEY_ISSUE_PARAGRAPH_1, AI_KEY_ISSUE_PARAGRAPH_2])
        self.assertNotEqual(payload["summary"]["key_issue_analysis"], payload["summary"]["key_issue_analysis_programmatic"])
        self.assertEqual(len(payload["summary"]["key_issue_items"]), 2)
        self.assertEqual(payload["summary"]["key_issue_items"][0]["heading"], "1. 血压控制现状与特征")
        self.assertEqual(payload["summary"]["key_issue_items"][1]["heading"], "2. 不良反应发生率")
        self.assertEqual(payload["summary"]["key_issue_items"][0]["chart_title"], "患者血压控制效果")
        self.assertEqual(payload["summary"]["key_issue_items"][1]["chart_title"], "不良反应发生率")
        self.assertEqual(payload["summary"]["key_issue_items"][0]["chart_type"], "pie")
        self.assertEqual(payload["summary"]["key_issue_items"][1]["chart_type"], "pie")
        self.assertGreaterEqual(len(payload["summary"]["overall_analysis"]), 2)
        self.assertLessEqual(sum(len(paragraph) for paragraph in payload["summary"]["overall_analysis"]), 700)
        self.assertEqual(payload["summary"]["overall_analysis"], payload["summary"]["overall_analysis_programmatic"])
        self.assertEqual(len(payload["summary"]["key_issue_analysis"]), 2)
        self.assertEqual(len(payload["summary"]["recommendations"]), 3)
        subtopic_lookup = {
            subtopic["subtitle"]: subtopic["paragraphs"]
            for section in payload["result_analysis"]["sections"]
            for subtopic in section["subtopics"]
        }
        for subtitle in ["头晕头痛不良反应分析", "购药渠道分析", "药品供应稳定性分析"]:
            paragraphs = subtopic_lookup[subtitle]
            self.assertEqual(len(paragraphs), 1)
            self.assertGreaterEqual(len(paragraphs[0]), MIN_ANALYSIS_CHARS)
            self.assertLessEqual(len(paragraphs[0]), MAX_ANALYSIS_CHARS)
            self.assertNotRegex(paragraphs[0], r"[ABCD]\.")
            self.assertRegex(paragraphs[0], r"[%％]")
            self.assertNotIn("选项A", paragraphs[0])
            self.assertNotIn("选项B", paragraphs[0])
            self.assertNotIn("逐项分布", paragraphs[0])
            self.assertTrue(any(marker in paragraphs[0] for marker in ["表明", "说明", "反映", "提示"]))
            self.assertTrue(any(marker in paragraphs[0] for marker in ["需", "仍", "可见", "整体", "进一步"]))

        analysis_paragraphs = [
            subtopic["paragraphs"][0]
            for section in payload["result_analysis"]["sections"]
            for subtopic in section["subtopics"]
        ]
        openings = [analysis_opening(paragraph) for paragraph in analysis_paragraphs]
        self.assertGreaterEqual(len(OPENING_STYLE_LIBRARY), 15)
        self.assertFalse(all(paragraph.startswith("从当前题目反馈分布看") for paragraph in analysis_paragraphs))
        self.assertGreaterEqual(len(set(openings)), int(len(openings) * 0.7))

    def test_adherence_payload_uses_diverse_analysis_openings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(minimal_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                adherence_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )

        analysis_paragraphs = [
            subtopic["paragraphs"][0]
            for section in payload["result_analysis"]["sections"]
            for subtopic in section["subtopics"]
        ]
        openings = [analysis_opening(paragraph) for paragraph in analysis_paragraphs]
        self.assertFalse(all(paragraph.startswith("从当前题目反馈分布看") for paragraph in analysis_paragraphs))
        self.assertGreaterEqual(len(set(openings)), int(len(openings) * 0.7))
        for paragraph in analysis_paragraphs:
            self.assertGreaterEqual(len(paragraph), MIN_ANALYSIS_CHARS)
            self.assertLessEqual(len(paragraph), MAX_ANALYSIS_CHARS)
            self.assertTrue(is_complete_analysis(paragraph))

    def test_build_payload_requires_ai_key_issue_analysis(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(minimal_markdown(include_key_issue=False), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            with self.assertRaisesRegex(ValueError, "5\\.1"):
                build_payload(
                    adherence_questionnaire(),
                    meta,
                    content,
                    Namespace(
                        product=None,
                        region=None,
                        time=None,
                        attachment_name=None,
                        survey_period=None,
                        sample_size=None,
                        valid_count=None,
                        disclaimer_unit=None,
                    ),
                )

    def test_build_payload_rejects_invalid_ai_key_issue_analysis_length(self) -> None:
        bad_content = minimal_markdown().replace(AI_ADHERENCE_KEY_ISSUE_1, "过短的重点问题分析。")
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(bad_content, encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            with self.assertRaisesRegex(ValueError, "5\\.1"):
                build_payload(
                    adherence_questionnaire(),
                    meta,
                    content,
                    Namespace(
                        product=None,
                        region=None,
                        time=None,
                        attachment_name=None,
                        survey_period=None,
                        sample_size=None,
                        valid_count=None,
                        disclaimer_unit=None,
                    ),
                )

    def test_build_payload_rejects_conflicting_draft_structure(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(
                sample_markdown()
                .replace("### 4.1药品疗效", "### 4.8不存在的维度"),
                encoding="utf-8",
            )
            meta, content = parse_markdown_content(report_content)
            with self.assertRaisesRegex(ValueError, "AI draft"):
                build_payload(
                    efficacy_questionnaire(),
                    meta,
                    content,
                    Namespace(
                        product=None,
                        region=None,
                        time=None,
                        attachment_name=None,
                        survey_period=None,
                        sample_size=None,
                        valid_count=None,
                        disclaimer_unit=None,
                    ),
                )

    def test_normalize_survey_period_accepts_compact_chinese_range(self) -> None:
        self.assertEqual(
            normalize_survey_period("2025年9月1日至9月30日"),
            "2025年09月01日——2025年09月30日",
        )
        self.assertEqual(
            normalize_survey_period("2025年9月1日-10月2日"),
            "2025年09月01日——2025年10月02日",
        )

    def test_render_from_template_preserves_template_pages_and_fields(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )
            output_docx = Path(temp_dir) / "rendered.docx"
            TemplateRenderer(Path(payload["meta"]["template_doc"]), payload).render(output_docx)
            validate_docx(output_docx, payload)

            document = Document(output_docx)
            texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            all_text = "\n".join(texts)
            expected_table_questions = expected_result_table_questions(payload)
            for required in ["问卷调研服务结算", "目录", "前言", "项目背景", "项目开展情况", "问卷说明", "问卷结果分析", "调研结果", "免责申明"]:
                self.assertIn(required, texts)
            self.assertLess(texts.index("目录"), texts.index("前言"))
            self.assertIn(payload["report_title"], texts)
            self.assertIn("4.1 药品疗效", texts)
            self.assertIn("（1） 血压控制效果分析", texts)
            self.assertIn("（1） 服药频率满意度分析", texts)
            self.assertIn("（1） 价格负担影响分析", texts)
            self.assertIn("（1） 药品供应稳定性分析", texts)
            self.assertIn("（1） 说明书信息清晰度分析", texts)
            self.assertIn("附件2：问卷调查明细表", texts)
            for subtitle in ["（1） 头晕头痛不良反应分析", "（1） 购药渠道分析", "（1） 药品供应稳定性分析"]:
                subtitle_idx = texts.index(subtitle)
                body_text = texts[subtitle_idx + 1]
                self.assertGreaterEqual(len(body_text), MIN_ANALYSIS_CHARS)
                self.assertLessEqual(len(body_text), MAX_ANALYSIS_CHARS)
                self.assertNotRegex(body_text, r"[ABCD]\.")
                self.assertRegex(body_text, r"[%％]")
                self.assertNotIn("选项A", body_text)
                self.assertNotIn("逐项分布", body_text)
                self.assertTrue(any(marker in body_text for marker in ["表明", "说明", "反映", "提示"]))
            self.assertNotIn("E.E", all_text)
            self.assertNotRegex(all_text, r"\n#+")
            self.assertGreaterEqual(len(document.tables), len(expected_table_questions) + 1)
            self.assertEqual(len(document.sections), 5)
            self.assertEqual(document.sections[0].header.paragraphs[0].text, payload["header_text"])
            self.assertIn("项目工具：调查问卷，14道选择题。其内容涵盖药品疗效、药品安全性、用药行为与习惯、用药便利性、药品经济性、药品可及性、用药指导信息评价7大维度，全面覆盖患者用药全流程关键节点。", texts)
            self.assertIn("样本采集时间：2025年10月01日——2025年10月31日", texts)
            self.assertIn(f"服务单位：{payload['service']['unit']}", texts)
            self.assertIn(f"日期：{payload['service']['date']}", texts)
            self.assertLess(texts.index(f"服务单位：{payload['service']['unit']}"), texts.index(f"日期：{payload['service']['date']}"))
            self.assertIn(f"服务提供单位:{payload['service']['unit']}", texts)
            self.assertIn(payload["service"]["date"], texts)
            settlement = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
            self.assertEqual(settlement[1][3], "1789例")
            self.assertEqual(settlement[1][4], "178,900")
            self.assertEqual(settlement[2][4], "30,000")
            self.assertEqual(settlement[3][4], "208,900")
            service_para = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == f"服务单位：{payload['service']['unit']}")
            date_para = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == f"日期：{payload['service']['date']}")
            self.assertEqual(service_para.alignment, 2)
            self.assertEqual(date_para.alignment, 2)
            key_issue_idx = texts.index("5.1问卷重点问题分析")
            key_issue_end = texts.index("5.2调研结果分析")
            key_issue_body = texts[key_issue_idx + 1:key_issue_end]
            self.assertNotIn("重点问题分析", key_issue_body)
            self.assertNotIn("1. 血压控制现状与特征", key_issue_body)
            self.assertNotIn("2. 不良反应发生率", key_issue_body)
            self.assertEqual(key_issue_body, [AI_KEY_ISSUE_PARAGRAPH_1, AI_KEY_ISSUE_PARAGRAPH_2])
            self.assertNotIn("呈现出较明确的反馈集中趋势", "\n".join(key_issue_body))
            self.assertNotIn("该环节已经成为影响", "\n".join(key_issue_body))
            overall_idx = texts.index("5.2调研结果分析")
            overall_end = texts.index("5.3建议")
            overall_paragraphs = texts[overall_idx + 1:overall_end]
            self.assertGreaterEqual(len(overall_paragraphs), 2)
            self.assertLessEqual(sum(len(paragraph) for paragraph in overall_paragraphs), 700)
            first_key_issue_para = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == key_issue_body[0])
            self.assertEqual(first_key_issue_para.alignment, 3)
            self.assertEqual(first_key_issue_para.paragraph_format.line_spacing, 2.5)
            self.assertEqual(first_key_issue_para.paragraph_format.first_line_indent, 304800)
            disclaimer_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "免责申明")
            disclaimer_heading_pos = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "免责申明")
            disclaimer_item = next(
                paragraph
                for paragraph in document.paragraphs[disclaimer_heading_pos + 1:]
                if paragraph.text.strip().startswith("（1）")
            )
            disclaimer_unit = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == f"服务提供单位:{payload['service']['unit']}")
            self.assertEqual(disclaimer_heading.alignment, 1)
            self.assertEqual(disclaimer_item.alignment, 3)
            self.assertEqual(disclaimer_item.paragraph_format.first_line_indent, 0)
            self.assertEqual(disclaimer_unit.alignment, 2)
            attachment_question = next(paragraph for paragraph in document.paragraphs if "您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？" in paragraph.text.strip())
            attachment_option = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("A. 选项A"))
            self.assertEqual(attachment_question.runs[0].font.name, "宋体")
            self.assertEqual(attachment_option.runs[0].font.name, "宋体")

            analysis_table_questions = []
            for table in document.tables:
                first_cell_text = table.cell(0, 0).text.strip()
                normalized_text = re.sub(r"^\d+\.", "", first_cell_text)
                for question in expected_table_questions:
                    if normalized_text == question:
                        analysis_table_questions.append(normalized_text)
            self.assertEqual(len(analysis_table_questions), len(expected_table_questions))
            self.assertEqual(Counter(analysis_table_questions), Counter(expected_table_questions))
            self.assertIn("您对厄贝沙坦氢氯噻嗪片的服药频率满意度如何？", analysis_table_questions)
            self.assertIn("您认为厄贝沙坦氢氯噻嗪片的价格对您的用药负担影响如何？", analysis_table_questions)
            self.assertIn("您所在地区的厄贝沙坦氢氯噻嗪片供应是否稳定？", analysis_table_questions)
            self.assertIn("您对厄贝沙坦氢氯噻嗪片说明书中 “用法用量”“不良反应” 的标注清晰度评价如何？", analysis_table_questions)

            with ZipFile(output_docx) as zipped:
                xml = zipped.read("word/document.xml").decode("utf-8", "ignore")
                styles_xml = zipped.read("word/styles.xml").decode("utf-8", "ignore")
                theme_xml = zipped.read("word/theme/theme1.xml").decode("utf-8", "ignore")
                settings_xml = zipped.read("word/settings.xml").decode("utf-8", "ignore")
                chart_names = [name for name in zipped.namelist() if re.match(r"word/charts/chart\d+\.xml$", name)]
                media_files = [name for name in zipped.namelist() if name.startswith("word/media/") and name != "word/media/"]
            with ZipFile(Path(payload["meta"]["template_doc"])) as template_zipped:
                template_chart_names = [name for name in template_zipped.namelist() if re.match(r"word/charts/chart\d+\.xml$", name)]
                template_media_files = [name for name in template_zipped.namelist() if name.startswith("word/media/") and name != "word/media/"]
            self.assertIn('TOC \\o "1-3" \\h \\u', xml)
            self.assertNotIn("w:sdt", xml)
            self.assertIn("updateFields", settings_xml)
            self.assertIn('w:val="true"', settings_xml)
            self.assertIn("调研时间：2025年10月01日——2025年10月31日", xml)
            self.assertNotIn("2025年11月1日-11月30日", xml)
            self.assertIn("宋体", xml)
            self.assertIn("宋体", styles_xml)
            self.assertIn("宋体", theme_xml)
            self.assertNotIn("汉仪中宋简", xml + styles_xml + theme_xml)
            self.assertEqual(sorted(chart_names), sorted(template_chart_names + ["word/charts/chart3.xml", "word/charts/chart4.xml"]))
            self.assertIn("word/charts/chart3.xml", chart_names)
            self.assertIn("word/charts/chart4.xml", chart_names)
            self.assertGreaterEqual(len(media_files), len(template_media_files) + 2)

            preface_idx = texts.index("前言")
            title_idx = texts.index(payload["report_title"])
            preface_body = texts[preface_idx + 1:title_idx]
            self.assertEqual(len(preface_body), 2)
            self.assertIn(payload["preface"][0], preface_body)
            self.assertIn(payload["preface"][-1], preface_body)

            background_idx = texts.index("项目背景")
            execution_idx = texts.index("项目开展情况")
            background_body = texts[background_idx + 1:execution_idx]
            self.assertEqual(len(background_body), 2)
            self.assertEqual(background_body, payload["project_background"])

            drawing_indices = [i for i, paragraph in enumerate(document.paragraphs) if paragraph._element.xpath('.//w:drawing')]
            result_idx = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "问卷结果分析")
            section_41_idx = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "4.1 药品疗效")
            key_issue_heading_idx = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "5.1问卷重点问题分析")
            key_issue_end_idx = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "5.2调研结果分析")
            result_drawings = [i for i in drawing_indices if result_idx < i < section_41_idx]
            key_issue_drawings = [i for i in drawing_indices if key_issue_heading_idx < i < key_issue_end_idx]
            self.assertEqual(len(result_drawings), 2)
            self.assertEqual(len(key_issue_drawings), 2)
            key_issue_chart_refs = chart_refs_between(output_docx, "5.1问卷重点问题分析", "5.2调研结果分析")
            self.assertEqual(key_issue_chart_refs, ["charts/chart3.xml", "charts/chart4.xml"])
            self.assertEqual([chart_part_type(output_docx, ref) for ref in key_issue_chart_refs], ["pie3DChart", "pie3DChart"])

            xml_root = ET.fromstring(xml.encode("utf-8"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for paragraph in xml_root.findall(".//w:p", ns):
                para_text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
                if re.match(r"^（\\d+）", para_text):
                    ppr = paragraph.find("w:pPr", ns)
                    self.assertIsNone(ppr.find("w:numPr", ns) if ppr is not None else None)

            attachment1_idx = texts.index(f"附件1：{payload['attachments']['attachment1_name']}")
            attachment2_idx = texts.index("附件2：问卷调查明细表")
            attachment_body = texts[attachment1_idx + 1:attachment2_idx]
            expected_questions = payload["attachments"]["attachment1_questions"]
            self.assertEqual(attachment_body[0], "（1） 您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？")
            self.assertEqual(attachment_body[1:5], ["A. 选项A", "B. 选项B", "C. 选项C", "D. 选项D"])
            self.assertEqual(attachment_body[5], "（2） 您服用厄贝沙坦氢氯噻嗪片期间，是否出现过头晕、头痛症状？")
            self.assertNotIn("（1） 1.", "\n".join(attachment_body))
            self.assertNotIn("（2） 6.", "\n".join(attachment_body))
            for display_index, question in enumerate(expected_questions, start=1):
                formatted = f"（{display_index}） {re.sub(r'^\\s*\\d+\\s*[\\.．、]\\s*', '', question['question'])}"
                self.assertIn(formatted, attachment_body)
                question_pos = attachment_body.index(formatted)
                expected_options = [f"{opt['code']}. {opt['text']}" for opt in question["options"]]
                actual_options = attachment_body[question_pos + 1:question_pos + 1 + len(expected_options)]
                self.assertEqual(actual_options, expected_options)

    def test_render_from_template_does_not_fallback_to_first_visual_group_on_unmatched_table(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )

            tampered_template = Path(temp_dir) / "tampered-template.docx"
            with ZipFile(Path(payload["meta"]["template_doc"])) as src_zip, ZipFile(tampered_template, "w") as dst_zip:
                for item in src_zip.infolist():
                    data = src_zip.read(item.filename)
                    if item.filename == "word/document.xml":
                        text = data.decode("utf-8", "ignore").replace(
                            "1.您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？",
                            "X.您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？",
                            1,
                        )
                        data = text.encode("utf-8")
                    dst_zip.writestr(item, data)

            output_docx = Path(temp_dir) / "rendered.docx"
            TemplateRenderer(tampered_template, payload).render(output_docx)

            document = Document(output_docx)
            expected_table_questions = expected_result_table_questions(payload)
            analysis_table_questions = []
            for table in document.tables:
                first_cell_text = table.cell(0, 0).text.strip()
                for question in expected_table_questions:
                    if first_cell_text == question:
                        analysis_table_questions.append(first_cell_text)
            counts = Counter(analysis_table_questions)
            self.assertLessEqual(counts[payload["result_analysis"]["sections"][0]["visual_groups"][0]["table_data"]["question"]], 1)

    def test_render_from_template_recalculates_settlement_total(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size="984",
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )

            output_docx = Path(temp_dir) / "settlement.docx"
            TemplateRenderer(Path(payload["meta"]["template_doc"]), payload).render(output_docx)
            validate_docx(output_docx, payload)

            settlement = [[cell.text.strip() for cell in row.cells] for row in Document(output_docx).tables[0].rows]
            self.assertEqual(settlement[1][3], "984例")
            self.assertEqual(settlement[1][4], "98,400")
            self.assertEqual(settlement[2][4], "30,000")
            self.assertEqual(settlement[3][4], "128,400")

    def test_adherence_render_removes_extra_template_sections_and_passes_final_validation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(minimal_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                adherence_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )
            validate_payload(payload)
            self.assertEqual(payload["meta"]["template_type"], "依从性与用药习惯")
            self.assertEqual([section["section_number"] for section in payload["result_analysis"]["sections"]], ["4.1", "4.2", "4.3", "4.4"])

            output_docx = Path(temp_dir) / "adherence.docx"
            TemplateRenderer(Path(payload["meta"]["template_doc"]), payload).render(output_docx)
            validate_docx(output_docx, payload)

            texts = [paragraph.text.strip() for paragraph in Document(output_docx).paragraphs if paragraph.text.strip()]
            self.assertIn("4.1 药物认知与信息获取", texts)
            self.assertIn("4.4 健康教育与支持需求", texts)
            self.assertNotIn("4.5 药品经济性", texts)
            self.assertNotIn("4.6 药品可及性", texts)
            self.assertNotIn("4.7 用药指导信息评价", texts)

    def test_final_validator_rejects_historical_regressions(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )
            output_docx = Path(temp_dir) / "rendered.docx"
            TemplateRenderer(Path(payload["meta"]["template_doc"]), payload).render(output_docx)
            validate_docx(output_docx, payload)

            bad_payload = copy.deepcopy(payload)
            bad_payload["result_analysis"]["sections"] = bad_payload["result_analysis"]["sections"][:4]
            with self.assertRaisesRegex(FinalValidationError, "section (count|headings) mismatch"):
                validate_docx(output_docx, bad_payload)

            first_paragraph = payload["result_analysis"]["sections"][0]["subtopics"][0]["paragraphs"][0]
            old_style = (
                "建议继续优化。逐项分布来看，选择主要选项的患者占比达到39.13%，说明该题仍存在旧式百分比分析表达，"
                "并且把选项差异直接转写为数字判断，弱化了模板要求的主结论、解释和收束逻辑。整体来看，这类写法容易让报告回到机械枚举状态，"
                "无法体现患者真实反馈背后的行为含义，也不利于后续形成稳定的专业研判。建议继续沿用这种旧式表达会造成质量回退。"
            )
            old_style_docx = Path(temp_dir) / "old-style.docx"
            copy_docx_with_document_xml_replace(output_docx, old_style_docx, first_paragraph, old_style)
            with self.assertRaisesRegex(FinalValidationError, "Forbidden old-style analysis"):
                validate_docx(old_style_docx, payload)

            attachment_docx = Path(temp_dir) / "bad-attachment.docx"
            copy_docx_with_document_xml_replace(
                output_docx,
                attachment_docx,
                "（1） 您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？",
                "1.您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？",
            )
            with self.assertRaisesRegex(FinalValidationError, "raw numeric question prefixes|question order mismatch"):
                validate_docx(attachment_docx, payload)

            font_docx = Path(temp_dir) / "bad-font.docx"
            copy_docx_with_document_xml_replace(output_docx, font_docx, "宋体", "SimSun", count=1000)
            with self.assertRaisesRegex(FinalValidationError, "Unexpected font"):
                validate_docx(font_docx, payload)

    def test_derive_preface_and_background_follow_template_rules(self) -> None:
        preface = derive_preface("厄贝沙坦氢氯噻嗪片", "北京市", "1789")
        background = derive_project_background("厄贝沙坦氢氯噻嗪片", "北京市")

        self.assertEqual(len(preface), 2)
        self.assertEqual(len(background), 2)
        self.assertGreaterEqual(sum(len(p) for p in preface), 380)
        self.assertLessEqual(sum(len(p) for p in preface), 430)
        self.assertGreaterEqual(sum(len(p) for p in background), 180)
        self.assertLessEqual(sum(len(p) for p in background), 360)
        self.assertTrue(any("北京市" in p for p in preface))
        self.assertTrue(any("1789" in p for p in preface))
        self.assertTrue(any("北京市" in p for p in background))
        self.assertNotIn("通用高血压降压药", "".join(preface))
        self.assertNotIn("通用高血压降压药", "".join(background))

    def test_expression_analysis_fits_new_style_rules(self) -> None:
        options = [
            {"label": "A", "text": "血压稳定在正常范围", "count": "880", "pct": "49.19%"},
            {"label": "B", "text": "血压多数时间处于正常范围", "count": "804", "pct": "44.94%"},
            {"label": "C", "text": "血压偶尔超出正常范围", "count": "105", "pct": "5.87%"},
            {"label": "D", "text": "血压持续超出正常范围", "count": "0", "pct": "0.00%"},
        ]
        paragraph = build_analysis_paragraph("您服用厄贝沙坦氢氯噻嗪片后，血压控制效果如何？", options)
        self.assertGreaterEqual(len(paragraph), MIN_ANALYSIS_CHARS)
        self.assertLessEqual(len(paragraph), MAX_ANALYSIS_CHARS)
        self.assertNotRegex(paragraph, r"[ABCD]\.")
        self.assertRegex(paragraph, r"[%％]")
        self.assertNotIn("选项A", paragraph)
        self.assertNotIn("选项B", paragraph)
        self.assertNotIn("逐项分布", paragraph)
        self.assertNotIn("从共性特征看", paragraph)
        self.assertTrue(is_complete_analysis(paragraph))

    def test_expression_opening_style_library_rotates_stably(self) -> None:
        self.assertGreaterEqual(len(OPENING_STYLE_LIBRARY), 15)
        options = [
            {"label": "A", "text": "非常了解", "count": "410", "pct": "41.00%"},
            {"label": "B", "text": "比较了解", "count": "360", "pct": "36.00%"},
            {"label": "C", "text": "一般了解", "count": "180", "pct": "18.00%"},
            {"label": "D", "text": "不了解", "count": "50", "pct": "5.00%"},
        ]
        paragraphs = [
            build_analysis_paragraph("您对心达康胶囊作用的了解程度如何？", options, style_index=index)
            for index in range(15)
        ]
        openings = [analysis_opening(paragraph) for paragraph in paragraphs]
        self.assertEqual(len(set(openings)), 15)
        self.assertNotIn("从当前题目反馈分布看", openings)
        for paragraph in paragraphs:
            self.assertTrue(is_complete_analysis(paragraph))

    def test_derive_subtitle_uses_semantic_mapping_not_truncation(self) -> None:
        from scripts.cluster_dimensions import cluster_dimensions
        questionnaire = adherence_questionnaire()
        grouped = cluster_dimensions(questionnaire)
        for section in grouped["sections"]:
            for subtopic in section.get("subtopics", []):
                subtitle = subtopic.get("subtitle", "")
                for pattern in [r"^(您|你|是否|怎么|什么|多少|哪个)", r"分析$.*(?:您|是否|怎么)"]:
                    self.assertNotRegex(subtitle, pattern, f"口语化副标题: {subtitle}")

    def test_ai_dimensions_override_hardcoded_template(self) -> None:
        from scripts.cluster_dimensions import cluster_dimensions
        questionnaire = adherence_questionnaire()
        ai_dimensions = dynamic_ai_dimensions()
        grouped = cluster_dimensions(questionnaire, ai_dimensions=ai_dimensions)
        self.assertEqual(grouped["dimension_count"], 3)
        self.assertEqual(len(grouped["project_dimensions"]), 3)
        self.assertEqual(grouped["project_dimensions"][0], "药物认知与信息获取")
        self.assertEqual(grouped["project_dimensions"][1], "用药行为与依从性")
        self.assertEqual(grouped["project_dimensions"][2], "健康教育与支持需求")
        self.assertEqual(grouped["sections"][0]["section_number"], "4.1")
        self.assertEqual(grouped["sections"][0]["section_title"], "药物认知与信息获取")
        self.assertEqual(grouped["sections"][0]["section_intro"], "本维度用于观察患者对药物作用的了解程度及信息来源结构。")
        q01_refs = grouped["sections"][0]["subtopics"][0]["question_refs"]
        self.assertTrue(any("q01" in r for r in q01_refs))
        self.assertEqual(grouped["key_issue_preferred_sections"], ["4.1", "4.2"])
        has_chart = any(vg.get("chart_type") is not None for vg in grouped["sections"][0]["visual_groups"])
        self.assertTrue(has_chart)

    def test_ai_dimensions_build_payload_integration(self) -> None:
        ai_dimensions = dynamic_ai_dimensions()
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(minimal_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            from scripts.build_payload import parse_dimensions_from_meta
            parsed_dims = parse_dimensions_from_meta({"dimensions_json": json.dumps(ai_dimensions)})
            self.assertEqual(parsed_dims["dimensions"][0]["name"], "药物认知与信息获取")
            payload = build_payload(
                adherence_questionnaire(),
                {**meta, "dimensions_json": json.dumps(ai_dimensions)},
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )
            output_docx = Path(temp_dir) / "dynamic_dimensions.docx"
            TemplateRenderer(Path(payload["meta"]["template_doc"]), payload).render(output_docx)
            validate_docx(output_docx, payload)
        validate_payload(payload)
        self.assertEqual(len(payload["result_analysis"]["sections"]), 3)
        self.assertEqual(payload["result_analysis"]["sections"][0]["section_title"], "药物认知与信息获取")
        self.assertEqual(payload["meta"]["template_type"], "依从性与用药习惯")

    def test_dimensions_json_validation_rejects_bad_schema(self) -> None:
        from scripts.build_payload import parse_dimensions_from_meta

        invalid_cases = [
            ("not json", "dimensions_json must be valid JSON"),
            (json.dumps({"dimensions": []}), "dimensions_json.dimensions must be a non-empty list"),
            (json.dumps({"dimensions": [{"name": "", "intro": "说明", "subtopics": []}]}), "dimensions_json.dimensions[0].name is required"),
            (json.dumps({"dimensions": [{"name": "维度", "intro": "说明", "subtopics": [{"patterns": [], "subtitle": "标题"}]}]}), "patterns must be a non-empty list"),
            (json.dumps({"dimensions": [{"name": "维度", "intro": "说明", "subtopics": [{"patterns": ["["], "subtitle": "标题"}]}]}), "invalid regex"),
            (json.dumps({"dimensions": [{"name": "维度", "intro": "说明", "subtopics": [{"patterns": ["题目"], "subtitle": "标题"}], "charts": [{"patterns": ["题目"], "chart_type": "line", "chart_style_profile": "efficacy_pie"}]}]}), "chart_type must be one of"),
        ]

        for raw, message in invalid_cases:
            with self.subTest(message=message):
                with self.assertRaises(ValueError) as ctx:
                    parse_dimensions_from_meta({"dimensions_json": raw})
                self.assertIn(message, str(ctx.exception))

    def test_is_complete_analysis_rejects_old_style_or_invalid_length(self) -> None:
        old_style = "建议继续优化。就逐项分布而言，选择A.选项A的患者占39.13%。整体来看，前两项反馈占比很高。从共性特征看，当前反馈较积极。"
        too_short = "多数患者反馈较积极，但仍需继续观察。"
        too_long = "整体反馈积极，说明产品表现较稳。" * 30
        self.assertFalse(is_complete_analysis(old_style))
        self.assertFalse(is_complete_analysis(too_short))
        self.assertFalse(is_complete_analysis(too_long))

    def test_validate_payload_rejects_preface_background_that_break_rules(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )

        invalid_cases = [
            ("single preface", {"preface": ["仅有一段，内容虽然提到北京市患者调研和产品价值，但整体仍明显过短，无法满足模板化前言要求。"]}, "Preface must contain exactly 2 paragraphs"),
            ("single background", {"project_background": ["仅有一段背景描述，虽然提到北京市和调研必要性，但未按两段结构展开，故应直接失败。"]}, "Project background must contain exactly 2 paragraphs"),
            ("missing region", {"preface": payload["preface"][:], "project_background": [p.replace("北京市", "该地区") for p in payload["project_background"]]}, "Project background must mention the region"),
            ("duplicate sections", {"project_background": payload["preface"][:2]}, "Project background must not substantially repeat the preface"),
        ]

        for _, patch, message in invalid_cases:
            bad_payload = dict(payload)
            bad_payload.update(patch)
            with self.assertRaisesRegex(ValueError, message):
                validate_payload(bad_payload)

    def test_validate_payload_rejects_invalid_overview_chart_types_or_values(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload(
                efficacy_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    attachment_name=None,
                    survey_period=None,
                    sample_size=None,
                    valid_count=None,
                    disclaimer_unit=None,
                ),
            )

        payload["result_analysis"]["overview_charts"][1]["chart_type"] = "pie"
        with self.assertRaisesRegex(ValueError, "must be bar"):
            validate_payload(payload)

        payload["result_analysis"]["overview_charts"][1]["chart_type"] = "bar"
        payload["result_analysis"]["overview_charts"][1]["render_mode"] = "office"
        with self.assertRaisesRegex(ValueError, "must use image render mode"):
            validate_payload(payload)

        payload["result_analysis"]["overview_charts"][1]["render_mode"] = "image"
        payload["result_analysis"]["overview_charts"][0]["render_mode"] = "office"
        with self.assertRaisesRegex(ValueError, "must use image render mode"):
            validate_payload(payload)

        payload["result_analysis"]["overview_charts"][0]["render_mode"] = "image"
        payload["result_analysis"]["overview_charts"][1]["values"] = [1.5, 2, 3]
        with self.assertRaisesRegex(ValueError, "integer question counts"):
            validate_payload(payload)


if __name__ == "__main__":
    unittest.main()
