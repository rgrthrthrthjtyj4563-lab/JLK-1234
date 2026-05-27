#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document

try:
    from .expression_data import MAX_ANALYSIS_CHARS, MIN_ANALYSIS_CHARS
except ImportError:
    from expression_data import MAX_ANALYSIS_CHARS, MIN_ANALYSIS_CHARS


EXPECTED_FONT = "宋体"
FORBIDDEN_ANALYSIS_PATTERNS = [
    r"\b[ABCD]\.",
    r"选项[ABCD]",
    r"逐项分布",
    r"从共性特征看",
    r"建议",
]
MIN_KEY_ISSUE_CHARS = 250
MAX_KEY_ISSUE_CHARS = 350
FORBIDDEN_KEY_ISSUE_PATTERNS = [
    r"呈现出较明确的反馈集中趋势",
    r"该环节已经成为影响",
    r"更适合作为后续患者教育和随访沟通的重点切入点",
]


class FinalValidationError(ValueError):
    pass


def _clean_text(text: str) -> str:
    return re.sub(r"[\s\u200b]+", "", str(text or ""))


def _visible_paragraph_texts(doc: Document) -> list[str]:
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]


def _find_text_index(texts: list[str], target: str, start: int = 0) -> int:
    for index in range(start, len(texts)):
        if texts[index] == target:
            return index
    raise FinalValidationError(f"Missing required paragraph: {target}")

def _section_heading_text(section: dict) -> str:
    return f"{section.get('section_number', '').strip()} {section.get('section_title', '').strip()}".strip()


def _subtopic_heading_text(index: int, subtopic: dict) -> str:
    return f"（{index}） {subtopic.get('subtitle', '').strip()}".strip()



def _validate_result_sections(texts: list[str], payload: dict) -> None:
    result_start = _find_text_index(texts, "问卷结果分析")
    try:
        result_end = _find_text_index(texts, "调研结果", result_start + 1)
    except FinalValidationError:
        result_end = _find_text_index(texts, "5.1问卷重点问题分析", result_start + 1)

    expected_section_count = len(payload.get("result_analysis", {}).get("sections", []))
    actual = [
        text
        for text in texts[result_start + 1:result_end]
        if re.match(r"^4\.\d+\S*", text)
    ]
    if len(actual) != expected_section_count:
        raise FinalValidationError(
            f"Result-analysis section count mismatch. expected={expected_section_count}, actual={len(actual)}"
        )
    expected = [
        _section_heading_text(section)
        for section in payload["result_analysis"]["sections"]
    ]
    if actual != expected:
        raise FinalValidationError(f"Result-analysis section headings mismatch. expected={expected}, actual={actual}")


def _validate_analysis_paragraphs(texts: list[str], payload: dict) -> None:
    result_start = _find_text_index(texts, "问卷结果分析")
    result_end = _find_text_index(texts, "调研结果", result_start + 1)
    for text in texts[result_start + 1:result_end]:
        for pattern in FORBIDDEN_ANALYSIS_PATTERNS:
            if re.search(pattern, text):
                raise FinalValidationError("Forbidden old-style analysis text in result-analysis section.")

    for section in payload["result_analysis"]["sections"]:
        for subtopic in section.get("subtopics", []):
            for paragraph in subtopic.get("paragraphs", []):
                if paragraph not in texts:
                    raise FinalValidationError(f"Missing result-analysis paragraph for subtitle: {subtopic.get('subtitle', '')}")
                for pattern in FORBIDDEN_ANALYSIS_PATTERNS:
                    if re.search(pattern, paragraph):
                        raise FinalValidationError(f"Forbidden old-style analysis text after subtitle: {subtopic.get('subtitle', '')}")
                if len(paragraph) < MIN_ANALYSIS_CHARS or len(paragraph) > MAX_ANALYSIS_CHARS:
                    raise FinalValidationError(f"Analysis paragraph length invalid after subtitle: {subtopic.get('subtitle', '')}")


def _analysis_opening(text: str) -> str:
    normalized = str(text or "").strip()
    match = re.match(r"^([^，。；;,.]{2,32})[，。；;,.]", normalized)
    if match:
        return match.group(1)
    return normalized[:16]


def _validate_analysis_opening_diversity(payload: dict) -> None:
    paragraphs = [
        paragraph
        for section in payload.get("result_analysis", {}).get("sections", [])
        for subtopic in section.get("subtopics", [])
        for paragraph in subtopic.get("paragraphs", [])
        if str(paragraph or "").strip()
    ]
    if len(paragraphs) < 4:
        return

    openings = [_analysis_opening(paragraph) for paragraph in paragraphs]
    if all(paragraph.startswith("从当前题目反馈分布看") for paragraph in paragraphs):
        raise FinalValidationError("Result-analysis paragraphs reuse the same legacy opening phrase.")

    max_count = max(openings.count(opening) for opening in set(openings))
    if max_count / len(openings) > 0.5:
        raise FinalValidationError("Result-analysis paragraph openings are too repetitive.")


def _attachment_question_text(display_index: int, question_text: str) -> str:
    cleaned = re.sub(r"^\s*\d+\s*[\.．、]\s*", "", str(question_text).strip())
    return f"（{display_index}） {cleaned}"


def _validate_attachment1(texts: list[str], payload: dict) -> None:
    attachment_name = payload.get("attachments", {}).get("attachment1_name", "问卷调查附件")
    start = next((i for i, text in enumerate(texts) if text == f"附件1：{attachment_name}" or text.startswith("附件1：")), None)
    if start is None:
        raise FinalValidationError("Missing attachment 1 heading.")
    end = _find_text_index(texts, "附件2：问卷调查明细表", start + 1)
    body = texts[start + 1:end]
    if any(re.match(r"^\d+[\.．、]", text) for text in body):
        raise FinalValidationError("Attachment 1 contains raw numeric question prefixes.")

    position = 0
    for display_index, question in enumerate(payload["attachments"]["attachment1_questions"], start=1):
        expected_question = _attachment_question_text(display_index, question.get("question", ""))
        if position >= len(body) or body[position] != expected_question:
            raise FinalValidationError(f"Attachment 1 question order mismatch at item {display_index}.")
        position += 1
        expected_options = [f"{option.get('code', '')}. {option.get('text', '')}" for option in question.get("options", [])]
        actual_options = body[position:position + len(expected_options)]
        if actual_options != expected_options:
            raise FinalValidationError(f"Attachment 1 option order mismatch at item {display_index}.")
        position += len(expected_options)


def _expected_result_table_questions(payload: dict) -> list[str]:
    qmap = {
        question.get("question_ref"): question.get("question", "")
        for question in payload.get("attachments", {}).get("attachment1_questions", [])
    }
    expected: list[str] = []
    seen: set[str] = set()
    for section in payload.get("result_analysis", {}).get("sections", []):
        for subtopic in section.get("subtopics", []):
            for question_ref in subtopic.get("question_refs", []):
                question = qmap.get(question_ref)
                if question and question_ref not in seen:
                    expected.append(question)
                    seen.add(question_ref)
    return expected


def _validate_result_tables(doc: Document, payload: dict) -> None:
    expected = [_clean_text(question) for question in _expected_result_table_questions(payload)]
    allowed = {
        _clean_text(question.get("question", ""))
        for question in payload.get("attachments", {}).get("attachment1_questions", [])
    }
    actual: list[str] = []
    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        first_cell = _clean_text(table.cell(0, 0).text)
        if first_cell in allowed:
            actual.append(first_cell)

    if actual != expected:
        raise FinalValidationError(
            f"Result-analysis table questions mismatch. expected={expected}, actual={actual}"
        )


def _body_items_between_result_and_summary(docx_path: Path) -> list[tuple[str, str]]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zipped:
        root = ET.fromstring(zipped.read("word/document.xml"))
    body = root.find("w:body", ns)
    if body is None:
        return []

    def text_of(element) -> str:
        return "".join(node.text or "" for node in element.findall(".//w:t", ns)).strip()

    raw_items: list[tuple[str, str]] = []
    for child in list(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = text_of(child)
            if text:
                raw_items.append(("p", text))
        elif tag == "tbl":
            raw_items.append(("tbl", text_of(child)))

    start = next((idx for idx, item in enumerate(raw_items) if item == ("p", "问卷结果分析")), None)
    end = next((idx for idx, item in enumerate(raw_items) if idx > (start or 0) and item == ("p", "调研结果")), None)
    if start is None or end is None:
        raise FinalValidationError("Cannot locate result-analysis XML range.")
    return raw_items[start + 1:end]


def _validate_result_table_blocks(docx_path: Path, payload: dict) -> None:
    qmap = {
        question.get("question_ref"): question.get("question", "")
        for question in payload.get("attachments", {}).get("attachment1_questions", [])
    }
    items = _body_items_between_result_and_summary(docx_path)
    cursor = 0
    first_section = next(iter(payload.get("result_analysis", {}).get("sections", [])), None)
    if first_section:
        first_heading = _section_heading_text(first_section)
        while cursor < len(items) and items[cursor] != ("p", first_heading):
            cursor += 1

    for section in payload.get("result_analysis", {}).get("sections", []):
        heading = _section_heading_text(section)
        if cursor >= len(items) or items[cursor] != ("p", heading):
            raise FinalValidationError(f"Missing result-analysis section block: {heading}")
        cursor += 1

        for subtopic_index, subtopic in enumerate(section.get("subtopics", []), start=1):
            subtitle = _clean_text(_subtopic_heading_text(subtopic_index, subtopic))
            if cursor >= len(items) or _clean_text(items[cursor][1]) != subtitle:
                raise FinalValidationError(f"Missing result-analysis subtopic block: {subtitle}")
            cursor += 1

            expected_questions = [_clean_text(qmap.get(ref, "")) for ref in subtopic.get("question_refs", []) if qmap.get(ref)]
            actual_questions: list[str] = []
            while cursor < len(items) and items[cursor][0] == "tbl":
                table_text = _clean_text(items[cursor][1])
                matched = next((question for question in expected_questions if table_text.startswith(question)), "")
                if matched:
                    actual_questions.append(matched)
                cursor += 1
            if actual_questions != expected_questions:
                raise FinalValidationError(
                    f"Result-analysis table block mismatch after {subtitle}. "
                    f"expected={expected_questions}, actual={actual_questions}"
                )
            if not actual_questions:
                raise FinalValidationError(f"Missing result-analysis table after subtopic: {subtitle}")

            expected_paragraphs = [_clean_text(text) for text in subtopic.get("paragraphs", [])]
            if not expected_paragraphs or cursor >= len(items) or _clean_text(items[cursor][1]) != expected_paragraphs[0]:
                raise FinalValidationError(f"Missing analysis paragraph after tables for subtopic: {subtitle}")
            cursor += 1


def _document_xml_root(docx_path: Path):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zipped:
        return ET.fromstring(zipped.read("word/document.xml")), ns


def _paragraph_text(element, ns: dict[str, str]) -> str:
    return "".join(node.text or "" for node in element.findall(".//w:t", ns)).strip()


def _validate_sample_size_text(texts: list[str], payload: dict) -> None:
    full_text = "\n".join(texts)
    if "None份" in full_text or "None名" in full_text or "nan份" in full_text:
        raise FinalValidationError("Rendered report contains invalid None/nan sample size text.")
    sample_size = str(payload.get("meta", {}).get("sample_size") or "").strip()
    if sample_size and f"{sample_size}份" not in full_text:
        raise FinalValidationError(f"Rendered report is missing sample size text: {sample_size}份")


def _format_money(value: object) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return str(value or "")


def _validate_settlement_table(doc: Document, payload: dict) -> None:
    settlement = payload.get("service", {}).get("settlement", {})
    if not settlement:
        return
    if not doc.tables:
        raise FinalValidationError("Missing settlement table.")
    table = doc.tables[0]
    if len(table.rows) < 4:
        raise FinalValidationError("Settlement table is incomplete.")
    expected_sample_amount = int(settlement.get("sample_count", 0)) * int(settlement.get("sample_unit_price", 0))
    expected_report_amount = int(settlement.get("report_count", 0)) * int(settlement.get("report_unit_price", 0))
    expected_total = expected_sample_amount + expected_report_amount
    expected = {
        (1, 3): f"{settlement.get('sample_count', 0)}例",
        (1, 4): _format_money(expected_sample_amount),
        (2, 4): _format_money(expected_report_amount),
        (3, 4): _format_money(expected_total),
    }
    for (row_idx, col_idx), value in expected.items():
        actual = table.cell(row_idx, col_idx).text.strip()
        if actual != value:
            raise FinalValidationError(f"Settlement table mismatch at row={row_idx}, col={col_idx}: expected={value}, actual={actual}")


def _validate_subtopic_numbering_xml(docx_path: Path) -> None:
    root, ns = _document_xml_root(docx_path)
    inside_ch4 = False
    for paragraph in root.findall(".//w:p", ns):
        text = _paragraph_text(paragraph, ns)
        if text == "问卷结果分析":
            inside_ch4 = True
        elif text == "调研结果":
            inside_ch4 = False
        if inside_ch4 and re.match(r"^（\d+）", text):
            if paragraph.find("./w:pPr/w:numPr", ns) is not None:
                raise FinalValidationError(f"Result-analysis subtopic keeps Word numbering: {text}")


def _validate_attachment_numbering_xml(docx_path: Path) -> None:
    root, ns = _document_xml_root(docx_path)
    inside_attachment = False
    for paragraph in root.findall(".//w:p", ns):
        text = _paragraph_text(paragraph, ns)
        if text.startswith("附件1"):
            inside_attachment = True
            continue
        if text.startswith("附件2"):
            inside_attachment = False
        if inside_attachment and text:
            if paragraph.find("./w:pPr/w:numPr", ns) is not None:
                raise FinalValidationError(f"Attachment 1 paragraph keeps Word numbering: {text}")


def _validate_png_chart_layout(docx_path: Path, payload: dict) -> None:
    root, ns = _document_xml_root(docx_path)
    paragraphs = root.findall(".//w:p", ns)
    texts = [_paragraph_text(paragraph, ns) for paragraph in paragraphs]
    chart_ns = {"c": "http://schemas.openxmlformats.org/drawingml/2006/chart"}

    def index_of(target: str) -> int | None:
        for index, text in enumerate(texts):
            if text == target:
                return index
        return None

    result_idx = index_of("问卷结果分析")
    first_section = next(iter(payload.get("result_analysis", {}).get("sections", [])), {})
    section_idx = index_of(_section_heading_text(first_section)) if first_section else None
    key_start = index_of("5.1问卷重点问题分析")
    key_end = index_of("5.2调研结果分析")

    def drawings_between(start: int | None, end: int | None) -> tuple[int, list[str]]:
        if start is None or end is None or end <= start:
            return 0, []
        drawing_count = 0
        native_chart_refs = []
        for paragraph in paragraphs[start + 1:end]:
            if paragraph.findall(".//w:drawing", ns):
                drawing_count += 1
            for chart in paragraph.findall(".//c:chart", chart_ns):
                native_chart_refs.append(chart.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", ""))
        return drawing_count, native_chart_refs

    result_drawings, result_chart_refs = drawings_between(result_idx, section_idx)
    key_drawings, key_chart_refs = drawings_between(key_start, key_end)
    if result_drawings != 2 or len(result_chart_refs) != 0:
        raise FinalValidationError("Result-analysis overview charts must render as two PNG drawings, not native charts.")
    if payload.get("summary", {}).get("key_issue_items"):
        if key_drawings != 2 or len(key_chart_refs) != 2:
            raise FinalValidationError("5.1 key issue charts must render as two native editable charts.")

    with ZipFile(docx_path) as zipped:
        media = [name for name in zipped.namelist() if name.startswith("word/media/") and name != "word/media/"]
        chart_parts = [name for name in zipped.namelist() if re.match(r"word/charts/chart\d+\.xml$", name)]
        rels_root = ET.fromstring(zipped.read("word/_rels/document.xml.rels"))
        relmap = {rel.attrib.get("Id", ""): rel.attrib.get("Target", "") for rel in rels_root}
        key_chart_targets = [relmap.get(rel_id, "") for rel_id in key_chart_refs]
        if payload.get("summary", {}).get("key_issue_items"):
            if key_chart_targets != ["charts/chart3.xml", "charts/chart4.xml"]:
                raise FinalValidationError("5.1 native chart relationship targets are incorrect.")
            key_issue_items = payload.get("summary", {}).get("key_issue_items", [])[:2]
            for target, item in zip(key_chart_targets, key_issue_items):
                chart_xml = zipped.read(f"word/{target}")
                chart_root = ET.fromstring(chart_xml)
                if chart_root.find(".//c:pie3DChart", chart_ns) is None:
                    raise FinalValidationError("5.1 native chart is not a 3D pie chart.")
                title = "".join(node.text or "" for node in chart_root.findall(".//c:title//a:t", {
                    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
                    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                })).strip()
                if title != str(item.get("chart_title", "")).strip():
                    raise FinalValidationError("5.1 native chart title does not match payload.")
                categories = [node.text or "" for node in chart_root.findall(".//c:cat//c:strCache//c:pt/c:v", chart_ns)]
                if categories != [str(value) for value in item.get("categories", [])]:
                    raise FinalValidationError("5.1 native chart categories do not match payload.")
                values = [float(node.text or 0) for node in chart_root.findall(".//c:val//c:numCache//c:pt/c:v", chart_ns)]
                expected_values = [float(value) for value in item.get("values", [])]
                if expected_values and sum(expected_values) > 1.5:
                    expected_values = [value / 100 for value in expected_values]
                if len(values) != len(expected_values) or any(abs(left - right) > 0.000001 for left, right in zip(values, expected_values)):
                    raise FinalValidationError("5.1 native chart values do not match payload.")
    template_doc = payload.get("meta", {}).get("template_doc")
    if template_doc:
        with ZipFile(template_doc) as zipped:
            template_media = [name for name in zipped.namelist() if name.startswith("word/media/") and name != "word/media/"]
            template_chart_parts = [name for name in zipped.namelist() if re.match(r"word/charts/chart\d+\.xml$", name)]
        expected_chart_parts = sorted(template_chart_parts + ["word/charts/chart3.xml", "word/charts/chart4.xml"])
        if sorted(chart_parts) != expected_chart_parts:
            raise FinalValidationError("Rendered docx native chart part inventory is incorrect.")
        if len(media) < len(template_media) + 2:
            raise FinalValidationError("Rendered docx is missing generated PNG chart media.")


def _validate_key_issue_text(texts: list[str], payload: dict) -> None:
    start = _find_text_index(texts, "5.1问卷重点问题分析")
    end = _find_text_index(texts, "5.2调研结果分析", start + 1)
    body = texts[start + 1:end]
    old_headings = {
        str(item.get("heading", "")).strip()
        for item in payload.get("summary", {}).get("key_issue_items", [])
        if str(item.get("heading", "")).strip()
    }
    for text in body:
        stripped = text.strip()
        if stripped == "重点问题分析" or re.match(r"^\s*[12]\.\s*重点问题分析\s*$", stripped):
            raise FinalValidationError("5.1 contains generic fallback title: 重点问题分析")
        if stripped in old_headings:
            raise FinalValidationError("5.1 contains numbered issue title paragraphs.")
        for pattern in FORBIDDEN_KEY_ISSUE_PATTERNS:
            if re.search(pattern, stripped):
                raise FinalValidationError("5.1 contains fixed programmatic wording.")
    expected = [str(item).strip() for item in payload.get("summary", {}).get("key_issue_analysis", []) if str(item).strip()]
    if body != expected:
        raise FinalValidationError("5.1 key issue text does not match AI payload paragraphs.")
    if len(body) != len(payload.get("summary", {}).get("key_issue_items", [])):
        raise FinalValidationError("5.1 paragraph count does not match key issue item count.")
    for index, paragraph in enumerate(body, start=1):
        if len(paragraph) < MIN_KEY_ISSUE_CHARS or len(paragraph) > MAX_KEY_ISSUE_CHARS:
            raise FinalValidationError(f"5.1 paragraph {index} length is invalid.")


def _validate_font_xml(docx_path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zipped:
        root = ET.fromstring(zipped.read("word/document.xml"))
        styles = ET.fromstring(zipped.read("word/styles.xml"))
        theme_text = zipped.read("word/theme/theme1.xml").decode("utf-8", "ignore")

    for run in root.findall(".//w:r", ns):
        texts = [node.text or "" for node in run.findall("w:t", ns)]
        if not "".join(texts).strip():
            continue
        rpr = run.find("w:rPr", ns)
        fonts = rpr.find("w:rFonts", ns) if rpr is not None else None
        if fonts is None:
            raise FinalValidationError("A visible text run is missing explicit font settings.")
        values = [
            fonts.get(f"{{{ns['w']}}}ascii"),
            fonts.get(f"{{{ns['w']}}}hAnsi"),
            fonts.get(f"{{{ns['w']}}}eastAsia"),
        ]
        if any(value != EXPECTED_FONT for value in values):
            raise FinalValidationError(f"Unexpected font in visible text run: {values}")
    for fonts in styles.findall(".//w:rFonts", ns):
        values = [
            fonts.get(f"{{{ns['w']}}}ascii"),
            fonts.get(f"{{{ns['w']}}}hAnsi"),
            fonts.get(f"{{{ns['w']}}}eastAsia"),
            fonts.get(f"{{{ns['w']}}}cs"),
        ]
        if any(value and value != EXPECTED_FONT for value in values):
            raise FinalValidationError(f"Unexpected font in style definition: {values}")
    for old in ["汉仪中宋简", "Times New Roman", "黑体", "Arial"]:
        if old in theme_text:
            raise FinalValidationError(f"Unexpected theme font remains: {old}")


def _validate_subtitle_formality(payload: dict) -> None:
    ORAL_INDICATORS = [
        r"^(您|你|我|他|她|它|咱们|大家)",
        r"是否",
        r"怎么会",
        r"什么$",
        r"多少",
        r"哪个",
        r"哪里",
        r"能不能",
        r"会不会",
        r"有没有",
    ]
    for section in payload.get("result_analysis", {}).get("sections", []):
        for subtopic in section.get("subtopics", []):
            subtitle = str(subtopic.get("subtitle", "")).strip()
            if not subtitle:
                continue
            for pattern in ORAL_INDICATORS:
                if re.search(pattern, subtitle):
                    raise FinalValidationError(
                        f"口语化副标题: '{subtitle}'，应为归纳性短语如'漏服应对行为分析'"
                    )


def validate_docx(docx_path: Path, payload: dict) -> None:
    doc = Document(str(docx_path))
    texts = _visible_paragraph_texts(doc)
    _validate_sample_size_text(texts, payload)
    _validate_settlement_table(doc, payload)
    _validate_result_sections(texts, payload)
    _validate_analysis_paragraphs(texts, payload)
    _validate_analysis_opening_diversity(payload)
    _validate_subtitle_formality(payload)
    _validate_result_tables(doc, payload)
    _validate_result_table_blocks(docx_path, payload)
    _validate_key_issue_text(texts, payload)
    _validate_attachment1(texts, payload)
    _validate_subtopic_numbering_xml(docx_path)
    _validate_attachment_numbering_xml(docx_path)
    _validate_png_chart_layout(docx_path, payload)
    _validate_font_xml(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate final rendered JLK patient report docx.")
    parser.add_argument("docx")
    parser.add_argument("payload_json")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))
    try:
        validate_docx(docx_path, payload)
    except FinalValidationError as exc:
        print(f"FINAL_VALIDATION_FAILED: {exc}", file=sys.stderr)
        raise SystemExit(1)


if __name__ == "__main__":
    main()
