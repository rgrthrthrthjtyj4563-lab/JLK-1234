#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from math import ceil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.colors as mcolors
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SONGTI = "汉仪中宋简"
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BLUE = "4684D3"
LIGHT_BLUE = "D5E4F5"


def pick_font_path() -> str | None:
    for path in [
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]:
        if Path(path).exists():
            return path
    return None


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def set_run_font(run, size_pt: float, bold: bool = False, east_asia: str = SONGTI) -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = BLACK


def set_header(doc: Document, text: str) -> None:
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if para.text or len(para.runs):
                p = para
                break
        else:
            p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 10.5, False)


def set_body_format(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 2.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_heading_format(paragraph, level: int) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    set_body_format(p)
    run = p.add_run(text)
    set_run_font(run, 12, False)


def add_heading(doc: Document, text: str, level: int) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    p = doc.add_paragraph(style=style_name)
    set_heading_format(p, level)
    size = 22 if level == 1 else 16 if level == 2 else 12
    bold = True if level < 3 else False
    run = p.add_run(text)
    set_run_font(run, size, bold)
    if level == 3:
        p.paragraph_format.first_line_indent = Pt(-10)
        p.paragraph_format.line_spacing = 2.5
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 22, True)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "右键更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    set_run_font(run, 14, False)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_settlement_table(doc: Document, sample_size: str) -> None:
    table = doc.add_table(rows=4, cols=5)
    table.style = "Normal Table"
    rows = [
        ["项目", "服务内容", "单价", "数量", "总额"],
        ["问卷样本", "抽取样本、数据清洗", "100元/例", f"{sample_size}例", f"{sample_size}00元"],
        ["问卷分析报告", "数据分析、可视化，形成市场分析报告", "30,000元/篇", "1篇", "30,000元"],
        ["合计", "", "", "", "详见结算单"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            if r == 0:
                set_cell_shading(cell, "4874CB")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, 11, r == 0)


def add_question_table(doc: Document, question: dict) -> None:
    options = question["options"]
    cols = len(options) + 2
    table = doc.add_table(rows=3, cols=cols)
    table.style = "Normal Table"
    row0 = [f"{question['number']}.{question['question']}", "选项"] + [f"{item['code']}.{item['text']}" for item in options]
    row1 = [f"{question['number']}.{question['question']}", "样本量"] + [item["count"] for item in options]
    row2 = [f"{question['number']}.{question['question']}", "占比"] + [item["pct"] for item in options]
    rows = [row0, row1, row2]
    fills = [HEADER_BLUE, LIGHT_BLUE, LIGHT_BLUE]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            set_cell_shading(cell, fills[r])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, 10.5, r == 0)


def gradient_bar(ax, bar, color_top: str, color_bottom: str) -> None:
    x, y = bar.get_x(), bar.get_y()
    w, h = bar.get_width(), bar.get_height()
    cmap = mcolors.LinearSegmentedColormap.from_list("", [color_bottom, color_top])
    grad = np.linspace(0, 1, 256).reshape(256, 1)
    im = ax.imshow(grad, extent=[x, x + w, y, y + h], aspect="auto", cmap=cmap, origin="lower", zorder=2)
    im.set_clip_path(bar)
    bar.set_facecolor((0, 0, 0, 0))
    bar.set_edgecolor("#5A87C8")


def create_pie_chart(path: Path, question: dict) -> None:
    font_path = pick_font_path()
    font = FontProperties(fname=font_path) if font_path else None
    labels = [f"{item['code']}.{item['text']}" for item in question["options"]]
    values = [float(item["pct"].rstrip("%")) for item in question["options"]]
    colors = ["#4F81BD", "#9BBB59", "#C0504D", "#8064A2", "#4BACC6", "#F79646"][: len(values)]
    fig, ax = plt.subplots(figsize=(6.8, 4.6), dpi=180)
    wedges, texts, autotexts = ax.pie(
        values,
        labels=labels,
        autopct="%1.2f%%",
        startangle=120,
        colors=colors,
        textprops={"fontproperties": font, "fontsize": 8},
        shadow=True,
        wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
    )
    for autotext in autotexts:
        autotext.set_fontproperties(font)
        autotext.set_fontsize(8)
        autotext.set_color("black")
    ax.axis("equal")
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_bar_chart(path: Path, question: dict) -> None:
    font_path = pick_font_path()
    font = FontProperties(fname=font_path) if font_path else None
    labels = [f"{item['code']}.{item['text']}" for item in question["options"]]
    values = [float(item["pct"].rstrip("%")) for item in question["options"]]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(7.2, 4.8), dpi=180)
    bars = ax.bar(x, values, color="none", width=0.55, zorder=3)
    for bar in bars:
        gradient_bar(ax, bar, "#8DB4E3", "#4F81BD")
    y_upper = max(60, int(ceil((max(values) + 5) / 10.0) * 10))
    ax.set_ylim(0, y_upper)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontproperties=font, fontsize=8)
    ticks = np.arange(0, y_upper + 1, 10)
    ax.set_yticks(ticks)
    ax.set_yticklabels([f"{int(tick)}%" for tick in ticks], fontproperties=font, fontsize=8)
    ax.grid(axis="y", color="#D9E2F3", linestyle="-", linewidth=0.7, zorder=0)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 1, f"{value:.2f}%", ha="center", va="bottom", fontproperties=font, fontsize=8)
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_chart(path: Path, question: dict, chart_type: str) -> None:
    if chart_type == "pie":
        create_pie_chart(path, question)
    else:
        create_bar_chart(path, question)


def render_report(payload: dict, output_docx: Path) -> Path:
    template_path = Path(payload["meta"]["template_doc"])
    doc = Document(str(template_path))
    clear_body(doc)
    set_header(doc, payload["header_text"])

    add_title(doc, "问卷调研服务结算")
    add_body(doc, "服务承诺：确保数据真实可靠，分析专业严谨，提供有价值的市场洞察，保持与客户的密切沟通，确保项目高质量完成。")
    p = doc.add_paragraph(style="Normal")
    set_body_format(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"\n服务单位：{payload['disclaimer']['unit']}\n日期：{payload['disclaimer']['date']}")
    set_run_font(run, 12, False)
    add_settlement_table(doc, payload["meta"]["sample_size"] or payload["meta"]["valid_count"] or "问卷未提供")
    add_page_break(doc)

    add_title(doc, "目录")
    add_toc_field(doc)
    add_page_break(doc)

    add_heading(doc, "前言", 1)
    for paragraph in payload["preface"]:
        add_body(doc, paragraph)
    add_page_break(doc)

    add_title(doc, payload["report_title"])
    add_heading(doc, "项目背景", 2)
    for paragraph in payload["project_background"]:
        add_body(doc, paragraph)

    add_heading(doc, "项目开展情况", 2)
    add_body(doc, payload["project_execution"]["overview"])
    add_body(doc, f"项目工具：{payload['project_execution']['tool']}")
    add_body(doc, f"调研维度：其内容涵盖{payload['project_execution']['dimensions']}，全面覆盖患者用药全流程关键节点。")
    add_body(doc, f"样本采集范围：{payload['project_execution']['scope']}")
    add_body(doc, f"样本采集数量：本次共收集筛选到有效问卷{payload['project_execution']['sample_count']}份。")
    add_body(doc, f"样本采集时间：{payload['project_execution']['period']}")

    add_heading(doc, "问卷说明", 2)
    add_body(doc, payload["questionnaire_note"]["intro"])
    for item in payload["questionnaire_note"]["items"]:
        add_body(doc, item)
    add_body(doc, payload["questionnaire_note"]["closing"])

    add_heading(doc, "问卷结果分析", 2)
    for paragraph in payload["result_analysis"]["intro"]:
        add_body(doc, paragraph)

    chart_dir = output_docx.parent / f"{output_docx.stem}_charts"
    chart_dir.mkdir(parents=True, exist_ok=True)
    for section in payload["result_analysis"]["sections"]:
        add_heading(doc, f"{section['section_number']}{section['section_title']}", 2)
        for intro in section["section_intro"]:
            add_body(doc, intro)

        visuals_by_ref = {item["question_ref"]: item for item in section["visual_groups"]}
        for subtopic in section["subtopics"]:
            add_heading(doc, subtopic["subtitle"], 3)
            for ref in subtopic["question_refs"]:
                if ref in visuals_by_ref:
                    visual = visuals_by_ref[ref]
                    add_question_table(doc, visual["table_data"])
                    if visual["chart_type"]:
                        chart_path = chart_dir / f"{visual['chart_ref']}.png"
                        create_chart(chart_path, visual["table_data"], visual["chart_type"])
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(str(chart_path), width=Cm(13.8))
                    break
            for paragraph in subtopic["paragraphs"]:
                add_body(doc, paragraph)

    add_heading(doc, "调研结果", 2)
    add_heading(doc, "5.1问卷重点问题分析", 2)
    for paragraph in payload["summary"]["key_issue_analysis"]:
        add_body(doc, paragraph)
    add_heading(doc, "5.2调研结果分析", 2)
    for paragraph in payload["summary"]["overall_analysis"]:
        add_body(doc, paragraph)
    add_heading(doc, "5.3建议", 2)
    for paragraph in payload["summary"]["recommendations"]:
        add_body(doc, paragraph)

    add_heading(doc, f"附件1：{payload['attachments']['attachment1_name']}", 1)
    for question in payload["attachments"]["attachment1_questions"]:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Pt(-10)
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(question["question"])
        set_run_font(run, 12, False)
        for option in question["options"]:
            op = doc.add_paragraph(style="Normal")
            op.paragraph_format.first_line_indent = Pt(12)
            op.paragraph_format.line_spacing = 2.0
            run = op.add_run(f"{option['code']}. {option['text']}")
            set_run_font(run, 12, False)

    add_heading(doc, f"附件2：{payload['attachments']['attachment2_name']}", 1)
    add_body(doc, f"01-1. 团队问卷调查明细-{payload['meta']['product']}-{payload['meta']['region']}")

    add_heading(doc, payload["disclaimer"]["title"], 1)
    for item in payload["disclaimer"]["items"]:
        add_body(doc, item)
    add_body(doc, f"服务提供单位:{payload['disclaimer']['unit']}")
    add_body(doc, payload["disclaimer"]["date"])

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def main() -> None:
    parser = argparse.ArgumentParser(description="Render template-aligned patient report docx from payload JSON.")
    parser.add_argument("payload_json")
    parser.add_argument("-o", "--output", required=True)
    args = parser.parse_args()

    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))
    output = Path(args.output)
    render_report(payload, output)
    print(output)


if __name__ == "__main__":
    main()
