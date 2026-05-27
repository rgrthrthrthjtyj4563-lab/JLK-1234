#!/usr/bin/env python3
"""
JLK-Pt-skill: Template-Driven Report Renderer (Object-Level Replacement)

STRATEGY: Do NOT clear the template body. Instead:
1. Open the template docx
2. Walk body elements, identify anchors by text/position
3. Replace: text in runs, data in table cells, chart data in XML
4. Preserve: all styles, section properties, headers/footers, drawings

This replaces the old render_report.py which used clear_body() and rebuilt
everything from scratch, losing all template formatting.
"""

from __future__ import annotations

import argparse
import copy
import json
import math
import re
import shutil
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from openpyxl import Workbook


# ─── OOXML Namespace constants ───────────────────────────────────────────────
NS = {
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "c14": "http://schemas.microsoft.com/office/drawing/2007/8/2/chart",
}
DEFAULT_FONT = "宋体"


def _register_namespaces():
    """Ensure consistent namespace prefixes in output XML."""
    for prefix, uri in NS.items():
        ET.register_namespace(prefix, uri)


_register_namespaces()


# ─── Helper: OOXML manipulation ──────────────────────────────────────────────

def _get_paragraph_text(p_element) -> str:
    """Extract plain text from a w:p element."""
    texts = []
    for r in p_element.findall(qn("w:r")):
        for t in r.findall(qn("w:t")):
            texts.append(t.text or "")
    return "".join(texts)


def _strip_numPr(p_element) -> None:
    """Remove <w:numPr> from a w:p element to suppress list numbering."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def _set_paragraph_text(p_element, text: str):
    """Replace all runs in a paragraph with a single run containing `text`."""
    # Remove all existing runs
    for r in list(p_element.findall(qn("w:r"))):
        p_element.remove(r)
    # Add new run
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    # Copy font properties from first run's rPr if available (keep styling)
    p_element.insert(0, r)


def _overwrite_paragraph_text_preserve_run_style(p_element, text: str):
    """Overwrite paragraph text while preserving the first run's styling."""
    runs = p_element.findall(qn("w:r"))
    text_runs = [r for r in runs if r.find(qn("w:drawing")) is None]
    if not text_runs:
        _set_paragraph_text(p_element, text)
        return

    first_run = text_runs[0]
    first_text = first_run.find(qn("w:t"))
    if first_text is None:
        first_text = OxmlElement("w:t")
        first_run.append(first_text)
    first_text.set(qn("xml:space"), "preserve")
    first_text.text = text

    for run in text_runs[1:]:
        for t in run.findall(qn("w:t")):
            t.text = ""


def _paragraph_has_drawing(p_element) -> bool:
    return any(
        run.find(qn("w:drawing")) is not None or run.find(qn("w:pict")) is not None
        for run in p_element.findall(qn("w:r"))
    )


def _paragraph_has_chart(p_element) -> bool:
    return bool(p_element.findall(".//c:chart", NS))


def _clone_first_chart_paragraph(p_element, relationship_id: str):
    cloned = copy.deepcopy(p_element)
    kept_chart = False
    for run in list(cloned.findall(qn("w:r"))):
        charts = run.findall(".//c:chart", NS)
        if not charts or kept_chart:
            cloned.remove(run)
            continue
        charts[0].set(qn("r:id"), relationship_id)
        for extra_chart in charts[1:]:
            parent = extra_chart.getparent() if hasattr(extra_chart, "getparent") else None
            if parent is not None:
                parent.remove(extra_chart)
        kept_chart = True
    return cloned if kept_chart else None


def _set_run_text(run_element, text: str):
    """Replace text in a specific w:r element."""
    for t in run_element.findall(qn("w:t")):
        t.text = text
        t.set(qn("xml:space"), "preserve")
        break
    else:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run_element.append(t)


def _replace_text_across_runs(p_element, old_text: str, new_text: str) -> bool:
    """Find `old_text` across runs and replace it with `new_text`. Returns True if replaced."""
    full_text = _get_paragraph_text(p_element)
    if old_text not in full_text:
        return False

    # Simple approach: join all text, do replacement, then distribute back
    runs = p_element.findall(qn("w:r"))
    if not runs:
        return False

    # Collect all text segments
    segments = []
    for r in runs:
        for t in r.findall(qn("w:t")):
            segments.append((r, t))

    # Build full text
    full = "".join(t.text or "" for _, t in segments)

    # Replace
    full = full.replace(old_text, new_text)

    # For simplicity, put everything in first run and clear others
    if segments:
        first_t = segments[0][1]
        first_t.text = full
        first_t.set(qn("xml:space"), "preserve")
        for _, t in segments[1:]:
            t.text = ""

    return True


def _replace_text_in_paragraph(p_element, old: str, new: str) -> bool:
    """Replace all occurrences of `old` with `new` in paragraph text."""
    full = _get_paragraph_text(p_element)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # Clear and rebuild
    for r in list(p_element.findall(qn("w:r"))):
        p_element.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_full
    r.append(t)
    p_element.append(r)
    return True


def _darken_hex(color: str, factor: float) -> str:
    """Return a darker #rrggbb color for simple pseudo-3D chart depth."""
    color = color.lstrip("#")
    if len(color) != 6:
        return "#666666"
    channels = [max(0, min(255, int(int(color[index:index + 2], 16) * factor))) for index in (0, 2, 4)]
    return "#" + "".join(f"{channel:02x}" for channel in channels)


def _ensure_paragraph_rpr(p_element):
    for r in p_element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        return rPr
    return None


def _apply_font_to_paragraph(p_element, font_name: str, size_pt: int, bold: bool = False):
    for r in p_element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)

        sz = rPr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rPr.append(sz)
        sz.set(qn("w:val"), str(size_pt * 2))

        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(size_pt * 2))

        bold_el = rPr.find(qn("w:b"))
        if bold and bold_el is None:
            bold_el = OxmlElement("w:b")
            rPr.append(bold_el)
        if bold_el is not None:
            bold_el.set(qn("w:val"), "1" if bold else "0")


def _apply_body_paragraph_layout(p_element, align: str = "both", first_line_chars: int = 200):
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)

    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), align)

    outline_lvl = pPr.find(qn("w:outlineLvl"))
    if outline_lvl is None:
        outline_lvl = OxmlElement("w:outlineLvl")
        pPr.append(outline_lvl)
    outline_lvl.set(qn("w:val"), "9")

    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line_chars > 0:
        ind.set(qn("w:firstLine"), str(int(first_line_chars * 2.4)))
        ind.set(qn("w:firstLineChars"), str(first_line_chars))
    else:
        ind.set(qn("w:firstLine"), "0")
        ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:rightChars"), "0")

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), "600")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    if qn("w:beforeLines") in spacing.attrib:
        del spacing.attrib[qn("w:beforeLines")]
    if qn("w:afterLines") in spacing.attrib:
        del spacing.attrib[qn("w:afterLines")]

    widow = pPr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        pPr.append(widow)
    widow.set(qn("w:val"), "0")

    text_direction = pPr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        pPr.append(text_direction)
    text_direction.set(qn("w:val"), "lrTb")

    snap = pPr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        pPr.append(snap)
    snap.set(qn("w:val"), "1")


def _set_paragraph_style_props(
    p_element,
    font_name: str,
    size_pt: int,
    bold: bool = False,
    align: str | None = None,
    body_layout: bool = False,
    first_line_chars: int = 200,
):
    _apply_font_to_paragraph(p_element, font_name, size_pt, bold)
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    if align:
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), align)
    if body_layout:
        _apply_body_paragraph_layout(p_element, align=align or "both", first_line_chars=first_line_chars)




# ─── Chart XML manipulation ──────────────────────────────────────────────────

def _update_chart_data(chart_xml_path: Path, categories: list[str], values: list[float], title: str = None) -> None:
    """Update the data in an Office-native chart XML file.

    Replaces <c:cat> category labels and <c:numCache> values,
    preserving all chart styling, 3D effects, gradients, etc.
    """
    tree = ET.parse(str(chart_xml_path))
    root = tree.getroot()

    c_ns = "http://schemas.openxmlformats.org/drawingml/2006/chart"
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

    # Update chart title if provided
    if title:
        for a_t in root.iter(f"{{{a_ns}}}t"):
            # Find the title text element (usually first a:t in c:title)
            parent = a_t
            for _ in range(10):  # walk up to find if we're in a c:title
                if parent is None:
                    break
                if parent.tag == f"{{{c_ns}}}title":
                    a_t.text = title
                    break
                parent = None  # Skip complex parent finding
            # Simpler: update all a:t that are likely titles
            # Actually, let's find it by structure
        # Find c:title > c:tx > c:rich > a:p > a:r > a:t
        for c_title in root.iter(f"{{{c_ns}}}title"):
            for a_t_el in c_title.iter(f"{{{a_ns}}}t"):
                a_t_el.text = title
                break

    # Update category labels in <c:cat>
    for c_cat in root.iter(f"{{{c_ns}}}cat"):
        # Clear existing strRef/numRef and strCache/numCache
        str_ref = c_cat.find(f"{{{c_ns}}}strRef")
        num_ref = c_cat.find(f"{{{c_ns}}}numRef")

        if str_ref is not None:
            formula = str_ref.find(f"{{{c_ns}}}f")
            if formula is not None:
                formula.text = f"Sheet1!$A$2:$A${len(categories) + 1}"
            # Update strCache
            str_cache = str_ref.find(f"{{{c_ns}}}strCache")
            if str_cache is not None:
                # Remove existing pt elements
                for pt in list(str_cache.findall(f"{{{c_ns}}}pt")):
                    str_cache.remove(pt)
                # Set count
                pt_count = str_cache.find(f"{{{c_ns}}}ptCount")
                if pt_count is None:
                    pt_count = ET.SubElement(str_cache, f"{{{c_ns}}}ptCount")
                pt_count.set("val", str(len(categories)))
                # Add new pts
                for idx, cat in enumerate(categories):
                    pt = ET.SubElement(str_cache, f"{{{c_ns}}}pt")
                    pt.set("idx", str(idx))
                    v = ET.SubElement(pt, f"{{{c_ns}}}v")
                    v.text = cat
        elif num_ref is not None:
            # Update numCache (categories as numbers)
            num_cache = num_ref.find(f"{{{c_ns}}}numCache")
            if num_cache is not None:
                for pt in list(num_cache.findall(f"{{{c_ns}}}pt")):
                    num_cache.remove(pt)
                pt_count = num_cache.find(f"{{{c_ns}}}ptCount")
                if pt_count is None:
                    pt_count = ET.SubElement(num_cache, f"{{{c_ns}}}ptCount")
                pt_count.set("val", str(len(categories)))
                for idx, cat in enumerate(categories):
                    pt = ET.SubElement(num_cache, f"{{{c_ns}}}pt")
                    pt.set("idx", str(idx))
                    v = ET.SubElement(pt, f"{{{c_ns}}}v")
                    v.text = cat

    # Update values in <c:numCache> inside <c:val>
    for c_val in root.iter(f"{{{c_ns}}}val"):
        num_ref = c_val.find(f"{{{c_ns}}}numRef")
        if num_ref is not None:
            formula = num_ref.find(f"{{{c_ns}}}f")
            if formula is not None:
                formula.text = f"Sheet1!$B$2:$B${len(values) + 1}"
            num_cache = num_ref.find(f"{{{c_ns}}}numCache")
            if num_cache is not None:
                format_code = num_cache.find(f"{{{c_ns}}}formatCode")
                for pt in list(num_cache.findall(f"{{{c_ns}}}pt")):
                    num_cache.remove(pt)
                pt_count = num_cache.find(f"{{{c_ns}}}ptCount")
                if pt_count is None:
                    pt_count = ET.SubElement(num_cache, f"{{{c_ns}}}ptCount")
                pt_count.set("val", str(len(values)))
                for idx, val in enumerate(values):
                    pt = ET.SubElement(num_cache, f"{{{c_ns}}}pt")
                    pt.set("idx", str(idx))
                    v = ET.SubElement(pt, f"{{{c_ns}}}v")
                    v.text = str(val)

    # Also try updating directly in <c:pie3DChart> or <c:barChart> series
    for c_ser in root.iter(f"{{{c_ns}}}ser"):
        cat_elem = c_ser.find(f"{{{c_ns}}}cat")
        val_elem = c_ser.find(f"{{{c_ns}}}val")
        if cat_elem is not None:
            str_ref = cat_elem.find(f"{{{c_ns}}}strRef")
            if str_ref is not None:
                formula = str_ref.find(f"{{{c_ns}}}f")
                if formula is not None:
                    formula.text = f"Sheet1!$A$2:$A${len(categories) + 1}"
                str_cache = str_ref.find(f"{{{c_ns}}}strCache")
                if str_cache is not None:
                    for pt in list(str_cache.findall(f"{{{c_ns}}}pt")):
                        str_cache.remove(pt)
                    pt_count = str_cache.find(f"{{{c_ns}}}ptCount")
                    if pt_count is None:
                        pt_count = ET.SubElement(str_cache, f"{{{c_ns}}}ptCount")
                    pt_count.set("val", str(len(categories)))
                    for idx, cat in enumerate(categories):
                        pt = ET.SubElement(str_cache, f"{{{c_ns}}}pt")
                        pt.set("idx", str(idx))
                        v = ET.SubElement(pt, f"{{{c_ns}}}v")
                        v.text = cat
        if val_elem is not None:
            num_ref = val_elem.find(f"{{{c_ns}}}numRef")
            if num_ref is not None:
                formula = num_ref.find(f"{{{c_ns}}}f")
                if formula is not None:
                    formula.text = f"Sheet1!$B$2:$B${len(values) + 1}"
                num_cache = num_ref.find(f"{{{c_ns}}}numCache")
                if num_cache is not None:
                    for pt in list(num_cache.findall(f"{{{c_ns}}}pt")):
                        num_cache.remove(pt)
                    pt_count = num_cache.find(f"{{{c_ns}}}ptCount")
                    if pt_count is None:
                        pt_count = ET.SubElement(num_cache, f"{{{c_ns}}}ptCount")
                    pt_count.set("val", str(len(values)))
                    for idx, val in enumerate(values):
                        pt = ET.SubElement(num_cache, f"{{{c_ns}}}pt")
                        pt.set("idx", str(idx))
                        v = ET.SubElement(pt, f"{{{c_ns}}}v")
                        v.text = str(val)

    tree.write(str(chart_xml_path), xml_declaration=True, encoding="UTF-8")


# ─── Main Renderer ───────────────────────────────────────────────────────────

class TemplateRenderer:
    """Opens a template docx and replaces content at known anchor points."""

    def __init__(self, template_path: Path, payload: dict):
        self.doc = Document(str(template_path))
        self.payload = payload
        self.body = self.doc._body._element
        self.meta = payload.get("meta", {})
        self._built_anchor_map()

    def _built_anchor_map(self):
        """Map body element indices to their text content for quick lookup."""
        self.anchors = {}
        self._body_children = []
        for i, child in enumerate(self.body):
            tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                break
            self._body_children.append(child)
            if tag == "p":
                text = _get_paragraph_text(child)
                self.anchors[i] = {"type": "p", "text": text, "elem": child}
            elif tag == "tbl":
                first_cell = child.find(f".//{qn('w:t')}")
                first_text = (first_cell.text or "")[:60] if first_cell is not None else ""
                self.anchors[i] = {"type": "tbl", "text": first_text, "elem": child}
            elif tag == "sdt":
                self.anchors[i] = {"type": "sdt", "elem": child}
            elif tag == "bookmarkEnd":
                self.anchors[i] = {"type": "bookmarkEnd", "elem": child}
            else:
                self.anchors[i] = {"type": tag, "elem": child}

    def _rebuild_anchors(self):
        """Rebuild anchor map after structural modifications."""
        self._built_anchor_map()

    def _find_paragraph(self, text_contains: str, start: int = 0) -> Optional[tuple[int, any]]:
        """Find the first paragraph element containing the given text."""
        for i, child in enumerate(self.body):
            if i < start:
                continue
            tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                break
            if tag == "p":
                full = _get_paragraph_text(child)
                if text_contains in full:
                    return i, child
        return None

    def _find_anchor_index(self, text_contains: str, start: int = 0) -> Optional[int]:
        """Find index of first element with text containing the given string."""
        for i, info in self.anchors.items():
            if i < start:
                continue
            if info["type"] == "p" and text_contains in info["text"]:
                return i
        return None

    def _find_exact_paragraph_index(self, text_exact: str, start: int = 0) -> Optional[int]:
        for i, info in self.anchors.items():
            if i < start:
                continue
            if info["type"] == "p" and info["text"].strip() == text_exact:
                return i
        return None

    def _get_paragraph_at(self, idx: int):
        """Get w:p element at body index."""
        child = self.body[idx]
        if child.tag.split("}")[1] != "p":
            raise ValueError(f"Element {idx} is not a paragraph")
        return child

    def _get_table_at(self, idx: int):
        """Get w:tbl element at body index."""
        child = self.body[idx]
        if child.tag.split("}")[1] != "tbl":
            raise ValueError(f"Element {idx} is not a table")
        return child

    # ─── Section 1: Header ───────────────────────────────────────────────

    def replace_header(self):
        """Replace header text in all sections."""
        new_header = self.payload.get("header_text", "")
        if not new_header:
            return
        for section in self.doc.sections:
            header = section.header
            for para in header.paragraphs:
                full = para.text
                if full.strip():
                    # Replace all text in runs
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_header
                    break

    # ─── Section 2: Settlement ──────────────────────────────────────────

    def replace_settlement(self):
        """Replace settlement table data."""
        # Find the table (first table in document)
        tbl_idx = None
        for i, info in self.anchors.items():
            if info["type"] == "tbl":
                tbl_idx = i
                break
        if tbl_idx is None:
            return

        tbl = self._get_table_at(tbl_idx)
        rows = tbl.findall(qn("w:tr"))

        settlement = self.payload.get("service", {}).get("settlement", {})
        sample_size = settlement.get("sample_count") or self.meta.get("sample_size") or self.meta.get("valid_count") or "问卷未提供"

        def money(value: object) -> str:
            try:
                return f"{int(value):,}"
            except Exception:
                return str(value or "")

        def set_cell_text(cell, text: str) -> None:
            paragraphs = cell.findall(qn("w:p"))
            if not paragraphs:
                p = OxmlElement("w:p")
                cell.append(p)
                paragraphs = [p]
            _set_paragraph_text(paragraphs[0], text)
            _set_paragraph_style_props(paragraphs[0], "宋体", 11, False, "center")
            for extra in paragraphs[1:]:
                cell.remove(extra)

        # Row 1: sample data (index 0-based = 1)
        if len(rows) >= 2:
            cells = rows[1].findall(qn("w:tc"))
            # cells[3] = count, cells[4] = total
            if len(cells) >= 4:
                set_cell_text(cells[3], f"{sample_size}例")
            if len(cells) >= 5:
                sample_amount = settlement.get("sample_amount")
                if sample_amount is None:
                    try:
                        sample_amount = int(sample_size) * 100
                    except Exception:
                        sample_amount = ""
                set_cell_text(cells[4], money(sample_amount))

        if len(rows) >= 3:
            cells = rows[2].findall(qn("w:tc"))
            if len(cells) >= 5:
                set_cell_text(cells[4], money(settlement.get("report_amount", 30000)))

        if len(rows) >= 4:
            cells = rows[3].findall(qn("w:tc"))
            if len(cells) >= 5:
                total = settlement.get("total_amount")
                if total is None:
                    try:
                        total = int(sample_size) * 100 + 30000
                    except Exception:
                        total = ""
                set_cell_text(cells[4], money(total))

    # ─── Section 3: Service Unit ────────────────────────────────────────

    def replace_service_unit(self):
        """Replace service unit name and date."""
        unit = self.payload.get("service", {}).get("unit", "")
        date_str = self.payload.get("service", {}).get("date", "")

        # Find paragraph after settlement table containing "服务单位"
        for i, info in self.anchors.items():
            if info["type"] == "p" and "服务单位" in info["text"]:
                p = self._get_paragraph_at(i)
                _set_paragraph_text(p, f"服务单位：{unit}")
                _set_paragraph_style_props(p, "宋体", 12, False, "right")
                if date_str:
                    new_p = OxmlElement("w:p")
                    r = OxmlElement("w:r")
                    t = OxmlElement("w:t")
                    t.set(qn("xml:space"), "preserve")
                    t.text = f"日期：{date_str}"
                    r.append(t)
                    new_p.append(r)
                    parent = self.body
                    parent.insert(list(parent).index(p) + 1, new_p)
                    _set_paragraph_style_props(new_p, "宋体", 12, False, "right")
                self._rebuild_anchors()
                break

    # ─── Section 4: TOC ────────────────────────────────────────────────

    def replace_toc(self):
        """Replace the SDT-based TOC with a fresh TOC field.

        The template has an SDT with pre-filled TOC entries pointing to specific
        _Toc bookmarks from the Guangdong data. We replace it with a clean TOC
        field that Word can regenerate.
        """
        # Find SDT element
        sdt_idx = None
        for i, info in self.anchors.items():
            if info["type"] == "sdt":
                sdt_idx = i
                break
        if sdt_idx is None:
            return

        # Replace SDT with a simple TOC field paragraph
        sdt_elem = self.body[sdt_idx]
        parent = self.body
        insert_at = list(parent).index(sdt_elem)

        # Preserve the visible "目录" heading. Some templates place the title
        # inside the same SDT block, so insert an explicit heading paragraph
        # when there is no standalone TOC title immediately before the SDT.
        has_toc_heading = False
        if insert_at > 0:
            prev_elem = parent[insert_at - 1]
            prev_tag = prev_elem.tag.split("}")[1] if "}" in prev_elem.tag else prev_elem.tag
            if prev_tag == "p" and _get_paragraph_text(prev_elem).strip() == "目录":
                has_toc_heading = True
        if not has_toc_heading:
            toc_heading = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = "目录"
            r.append(t)
            toc_heading.append(r)
            parent.insert(insert_at, toc_heading)
            _set_paragraph_style_props(toc_heading, "宋体", 16, True, "left")
            insert_at += 1

        # Create a new paragraph with TOC field
        new_p = OxmlElement("w:p")

        # Add TOC field
        r = OxmlElement("w:r")
        # Field begin
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r.append(fld_begin)
        new_p.append(r)

        # Field instruction
        r2 = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\u'
        r2.append(instr)
        new_p.append(r2)

        # Field separator
        r3 = OxmlElement("w:r")
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r3.append(fld_sep)
        new_p.append(r3)

        # Placeholder text
        r4 = OxmlElement("w:r")
        t4 = OxmlElement("w:t")
        t4.text = "（请在 Word 中右键此处 → 更新域，以生成目录）"
        r4.append(t4)
        new_p.append(r4)

        # Field end
        r5 = OxmlElement("w:r")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r5.append(fld_end)
        new_p.append(r5)

        # Replace SDT with new paragraph
        parent.replace(sdt_elem, new_p)
        _set_paragraph_style_props(new_p, "宋体", 12, False, "left")
        # Rebuild anchors because body structure changed
        self._rebuild_anchors()

    # ─── Section 5: Text replacements ───────────────────────────────────

    def _replace_all_text(self, old: str, new: str):
        """Replace all occurrences of `old` with `new` across ALL text nodes.
        
        Handles paragraphs, table cells, headers, and footnotes."""
        # 1. Body paragraphs and table cells
        for child in self.body:
            tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                break
            if tag == "p":
                self._safe_replace_in_paragraph(child, old, new)
            elif tag == "tbl":
                # Replace in all table cell paragraphs
                for tc in child.iter(qn("w:tc")):
                    for p in tc.findall(qn("w:p")):
                        self._safe_replace_in_paragraph(p, old, new)
        
        # 2. Headers and footers in all sections
        for section in self.doc.sections:
            for container in [section.header, section.footer,
                              section.first_page_header, section.first_page_footer]:
                try:
                    for p in container.paragraphs:
                        self._safe_replace_in_paragraph(p._element, old, new)
                except:
                    pass

    def _safe_replace_in_paragraph(self, p_element, old: str, new: str):
        """Replace text in a paragraph, preserving drawings."""
        has_drawing = any(
            r.find(qn("w:drawing")) is not None
            for r in p_element.findall(qn("w:r"))
        )
        if has_drawing:
            for r in p_element.findall(qn("w:r")):
                drawing = r.find(qn("w:drawing"))
                if drawing is not None:
                    continue
                for t in r.findall(qn("w:t")):
                    if t.text and old in t.text:
                        t.text = t.text.replace(old, new)
                        t.set(qn("xml:space"), "preserve")
        else:
            _replace_text_in_paragraph(p_element, old, new)

    def replace_metadata(self):
        """Replace all metadata placeholders."""
        product = self.meta.get("product", "")
        region = self.meta.get("region", "")
        sample_size = self.meta.get("sample_size") or self.meta.get("valid_count") or ""
        survey_period = self.meta.get("survey_period", "")
        service_date = self.payload.get("service", {}).get("date", "")

        # Replace "厄贝沙坦氢氯噻嗪片" with product name
        if product and product != "厄贝沙坦氢氯噻嗪片":
            self._replace_all_text("厄贝沙坦氢氯噻嗪片", product)

        # Replace "广东省" with region
        if region and region != "广东省":
            self._replace_all_text("广东省", region)

        # Replace sample count
        if sample_size:
            import re
            old_sample = re.search(r"\d+份", _get_paragraph_text(self.body[24])) if len(self.body) > 24 else None
            if old_sample:
                self._replace_all_text(old_sample.group(0), f"{sample_size}份")
            # Also replace the standalone number in "1642份" patterns
            self._replace_all_text("1642份", f"{sample_size}份")
            self._replace_all_text("1642名", f"{sample_size}名")

        # Replace survey period
        if survey_period:
            for i, info in self.anchors.items():
                if info["type"] == "p" and "调研时间" in info["text"]:
                    p = self._get_paragraph_at(i)
                    _overwrite_paragraph_text_preserve_run_style(p, f"调研时间：{survey_period}")
                    continue
                if info["type"] == "p" and "样本采集时间" in info["text"]:
                    p = self._get_paragraph_at(i)
                    _overwrite_paragraph_text_preserve_run_style(p, f"样本采集时间：{survey_period}")

        if service_date:
            self._replace_all_text("2025年12月11日", service_date)

    def replace_preface(self):
        """Replace preface text while preserving the template title anchor."""
        preface = self.payload.get("preface", [])
        if not preface:
            return

        idx = self._find_anchor_index("前言")
        if idx is None:
            return

        heading = self._get_paragraph_at(idx)
        _strip_numPr(heading)
        _set_paragraph_style_props(heading, "宋体", 22, True, "center")

        title_idx = self._find_anchor_index("调查问卷分析报告", start=idx + 1)
        if title_idx is None:
            title_idx = idx + 2

        body_paragraphs = []
        for i in range(idx + 1, title_idx):
            info = self.anchors.get(i)
            if info and info["type"] == "p" and _get_paragraph_text(info["elem"]).strip():
                body_paragraphs.append(info["elem"])

        if not body_paragraphs:
            return

        insert_before = self.anchors[title_idx]["elem"] if title_idx in self.anchors else body_paragraphs[-1]
        template_paragraph = body_paragraphs[-1]
        while len(body_paragraphs) < len(preface):
            cloned = copy.deepcopy(template_paragraph)
            insert_before.addprevious(cloned)
            body_paragraphs.append(cloned)

        for index, paragraph in enumerate(body_paragraphs):
            if index < len(preface):
                _set_paragraph_text(paragraph, preface[index])
                _set_paragraph_style_props(paragraph, "宋体", 12, False)
            else:
                parent = paragraph.getparent()
                if parent is not None:
                    parent.remove(paragraph)

        self._rebuild_anchors()

    def replace_report_title(self):
        """Replace the report title paragraph."""
        title = self.payload.get("report_title", "")
        if not title:
            return

        # Find the title paragraph (after 前言, before 项目背景)
        idx = self._find_anchor_index("用药体验与疗效反馈患者调查问卷分析报告")
        if idx is None:
            idx = self._find_anchor_index("调查问卷分析报告")
        if idx is None:
            return

        p = self._get_paragraph_at(idx)
        _set_paragraph_text(p, title)
        _set_paragraph_style_props(p, "宋体", 22, True, "center")

    def replace_project_background(self):
        """Replace project background paragraphs with structured 4-section content."""
        paragraphs = self.payload.get("project_background", [])
        if not paragraphs:
            return

        # Find "项目背景" heading
        idx = self._find_anchor_index("项目背景")
        if idx is None:
            return

        heading = self._get_paragraph_at(idx)
        _set_paragraph_style_props(heading, "宋体", 16, True, "left")

        # Background paragraphs follow the heading
        bg_indices = []
        for i in range(idx + 1, idx + 20):
            if i in self.anchors and self.anchors[i]["type"] == "p":
                text = self.anchors[i]["text"]
                if "项目开展情况" in text or "问卷说明" in text or text.strip().startswith("二、"):
                    break
                bg_indices.append(i)
            else:
                break

        # Replace with payload paragraphs
        for pi, bi in enumerate(bg_indices):
            if pi < len(paragraphs):
                p = self._get_paragraph_at(bi)
                _set_paragraph_text(p, paragraphs[pi])
                _set_paragraph_style_props(p, "宋体", 12, False)
            else:
                # Remove extra template paragraphs
                p = self._get_paragraph_at(bi)
                p.getparent().remove(p)

    def replace_project_execution(self):
        """Replace project execution paragraphs."""
        pe = self.payload.get("project_execution", {})
        if not pe:
            return

        lines = pe.get("lines", [])
        if not lines:
            return

        title_idx = self._find_anchor_index("项目开展情况")
        if title_idx is None:
            return
        heading = self._get_paragraph_at(title_idx)
        _set_paragraph_style_props(heading, "宋体", 16, True, "left")

        target_indices = []
        for i in range(title_idx + 1, title_idx + 12):
            if i not in self.anchors:
                break
            info = self.anchors[i]
            if info["type"] != "p":
                break
            text = info["text"].strip()
            if text in {"问卷说明", "三、问卷说明"}:
                break
            if text:
                target_indices.append(i)

        for line_index, line in enumerate(lines):
            if line_index < len(target_indices):
                p = self._get_paragraph_at(target_indices[line_index])
                _set_paragraph_text(p, line)
                _set_paragraph_style_props(p, "宋体", 12, False)
        for extra_idx in target_indices[len(lines):]:
            p = self._get_paragraph_at(extra_idx)
            p.getparent().remove(p)

    def replace_questionnaire_note(self):
        """Replace questionnaire note paragraphs."""
        qn_data = self.payload.get("questionnaire_note", {})
        if not qn_data:
            return

        # Find "问卷说明" heading
        idx = self._find_anchor_index("问卷说明")
        if idx is None:
            return
        heading = self._get_paragraph_at(idx)
        _set_paragraph_style_props(heading, "宋体", 16, True, "left")

        # Intro paragraph (first after heading)
        intro_idx = idx + 1
        if intro_idx in self.anchors:
            p = self._get_paragraph_at(intro_idx)
            _set_paragraph_text(p, qn_data.get("intro", ""))
            _set_paragraph_style_props(p, "宋体", 12, False)

        # Items: find the numbered items (1．, 2．, 3．, 4．)
        items = qn_data.get("items", [])
        item_indices = []
        for i in range(intro_idx + 1, intro_idx + 10):
            if i in self.anchors and self.anchors[i]["type"] == "p":
                text = self.anchors[i]["text"]
                if text.startswith(("1", "2", "3", "4")) and "．" in text[:3]:
                    item_indices.append(i)
                elif len(item_indices) >= len(items):
                    break

        for pi, bi in enumerate(item_indices):
            if pi < len(items):
                p = self._get_paragraph_at(bi)
                _set_paragraph_text(p, items[pi])
                _set_paragraph_style_props(p, "宋体", 12, False)

        # Closing paragraph (after items)
        closing_idx = (item_indices[-1] + 1) if item_indices else intro_idx + 5
        if closing_idx in self.anchors:
            p = self._get_paragraph_at(closing_idx)
            _set_paragraph_text(p, qn_data.get("closing", ""))
            _set_paragraph_style_props(p, "宋体", 12, False)

    def replace_result_analysis_intro(self):
        """Replace the intro paragraph for 问卷结果分析."""
        intro = self.payload.get("result_analysis", {}).get("intro", [])
        if not intro:
            return

        idx = self._find_anchor_index("问卷结果分析")
        if idx is None:
            return
        heading = self._get_paragraph_at(idx)
        _set_paragraph_style_props(heading, "宋体", 16, True, "left")

        # Intro is the paragraph right after (before the chart)
        intro_idx = idx + 1
        if intro_idx in self.anchors and self.anchors[intro_idx]["type"] == "p":
            p = self._get_paragraph_at(intro_idx)
            _set_paragraph_text(p, intro[0] if intro else "")
            _set_paragraph_style_props(p, "宋体", 12, False)

    def replace_analysis_sections(self):
        """Replace all 4.x analysis sections in-place.

        The result-analysis chapter is template-driven: keep the template's
        headings, tables, drawings, and paragraph styles, and only replace
        their content. If a payload section contains more subtopics than the
        template section has blocks, clone the last template block in that
        section and then replace the cloned content.
        """
        sections = self.payload.get("result_analysis", {}).get("sections", [])
        if not sections:
            return

        import re

        def local_tag(element) -> str:
            return element.tag.split("}")[1] if "}" in element.tag else element.tag

        def is_section_stop(element) -> bool:
            if local_tag(element) != "p":
                return False
            stripped = _get_paragraph_text(element).strip()
            return bool(
                re.match(r"^(4\.\d+|5\.\d+)", stripped)
                or stripped == "调研结果"
                or stripped.startswith("附件1")
                or stripped.startswith("附件2")
                or stripped == "免责申明"
            )

        def is_subtopic_heading(element) -> bool:
            if local_tag(element) != "p":
                return False
            stripped = _get_paragraph_text(element).strip()
            pPr = element.find(qn("w:pPr"))
            style_val = ""
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_val = pStyle.get(qn("w:val"), "")
            return style_val == "4" or "Heading" in style_val or bool(re.match(r"^[（(]\d+[)）]\s*", stripped))

        def remove_element(element) -> None:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

        def paragraph_has_visible_text(element) -> bool:
            return local_tag(element) == "p" and bool(_get_paragraph_text(element).strip())

        def collect_children():
            collected = []
            for child in self.body:
                if local_tag(child) == "sectPr":
                    break
                collected.append(child)
            return collected

        def build_units(section_elements):
            heading_positions = [
                idx for idx, element in enumerate(section_elements)
                if is_subtopic_heading(element)
            ]
            units = []
            for pos, heading_pos in enumerate(heading_positions):
                next_pos = heading_positions[pos + 1] if pos + 1 < len(heading_positions) else len(section_elements)
                unit_elements = section_elements[heading_pos:next_pos]
                units.append({"heading": section_elements[heading_pos], "elements": unit_elements})
            intro_elements = section_elements[:heading_positions[0]] if heading_positions else section_elements
            return intro_elements, units

        def clone_missing_units(units, required_count, insert_before):
            if not units or required_count <= len(units):
                return
            template_elements = units[-1]["elements"]
            for _ in range(required_count - len(units)):
                cloned_elements = [copy.deepcopy(element) for element in template_elements]
                for cloned in cloned_elements:
                    insert_before.addprevious(cloned)
                cloned_heading = next((element for element in cloned_elements if is_subtopic_heading(element)), cloned_elements[0])
                units.append({"heading": cloned_heading, "elements": cloned_elements})

        def ensure_tables(unit, required_count, insert_before):
            tables = [element for element in unit["elements"] if local_tag(element) == "tbl"]
            if tables and len(tables) < required_count:
                template_table = tables[-1]
                last_element = unit["elements"][-1]
                for _ in range(required_count - len(tables)):
                    cloned_table = copy.deepcopy(template_table)
                    last_element.addnext(cloned_table)
                    unit["elements"].append(cloned_table)
                    tables.append(cloned_table)
                    last_element = cloned_table
            return tables

        sec_map = {sec.get("section_number", ""): sec for sec in sections}
        children = collect_children()

        i = 0
        while i < len(children):
            child = children[i]
            if child.getparent() is None or local_tag(child) != "p":
                i += 1
                continue

            text = _get_paragraph_text(child).strip()
            match = re.match(r"^(4\.\d+)", text)
            if not match:
                i += 1
                continue

            sec_num = match.group(1)
            block_end = i + 1
            while block_end < len(children):
                if children[block_end].getparent() is not None and is_section_stop(children[block_end]):
                    break
                block_end += 1

            payload_sec = sec_map.get(sec_num)
            if payload_sec is None:
                for element in children[i:block_end]:
                    remove_element(element)
                i = block_end
                continue

            section_title = payload_sec.get("section_title", "")
            _set_paragraph_text(child, f"{sec_num} {section_title}".strip())
            _set_paragraph_style_props(child, "宋体", 16, True, "left")

            subtopics = payload_sec.get("subtopics", [])
            visual_groups = payload_sec.get("visual_groups", [])
            visual_by_ref = {
                visual.get("question_ref"): visual
                for visual in visual_groups
                if visual.get("question_ref")
            }

            section_elements = [element for element in children[i + 1:block_end] if element.getparent() is not None]
            intro_elements, units = build_units(section_elements)
            for element in intro_elements:
                if paragraph_has_visible_text(element) or local_tag(element) == "tbl":
                    remove_element(element)

            insert_before = children[block_end] if block_end < len(children) else None
            if insert_before is not None and insert_before.getparent() is not None:
                clone_missing_units(units, len(subtopics), insert_before)

            for unit_idx, unit in enumerate(units):
                if unit_idx >= len(subtopics):
                    for element in unit["elements"]:
                        remove_element(element)
                    continue

                subtopic = subtopics[unit_idx]
                subtitle = subtopic.get("subtitle", "")
                _set_paragraph_text(unit["heading"], f"（{unit_idx + 1}） {subtitle}".strip())
                _strip_numPr(unit["heading"])
                _set_paragraph_style_props(unit["heading"], "宋体", 12, True, "left")

                question_refs = subtopic.get("question_refs", [])
                tables = ensure_tables(unit, len(question_refs), unit["heading"])
                for table_idx, table in enumerate(tables):
                    if table_idx < len(question_refs):
                        visual = visual_by_ref.get(question_refs[table_idx])
                        if visual is not None:
                            self._replace_table_element(table, visual)
                        else:
                            remove_element(table)
                    else:
                        remove_element(table)

                body_paragraphs = [
                    element for element in unit["elements"]
                    if local_tag(element) == "p"
                    and element is not unit["heading"]
                    and not _paragraph_has_drawing(element)
                    and paragraph_has_visible_text(element)
                    and element.getparent() is not None
                ]
                empty_paragraphs = [
                    element for element in unit["elements"]
                    if local_tag(element) == "p"
                    and element is not unit["heading"]
                    and not _paragraph_has_drawing(element)
                    and not paragraph_has_visible_text(element)
                    and element.getparent() is not None
                ]
                for empty_paragraph in empty_paragraphs:
                    remove_element(empty_paragraph)

                paragraphs = [text for text in subtopic.get("paragraphs", []) if text and text.strip()]
                for para_idx, paragraph in enumerate(body_paragraphs):
                    if para_idx < len(paragraphs):
                        _set_paragraph_text(paragraph, paragraphs[para_idx])
                        _set_paragraph_style_props(paragraph, "宋体", 12, False)
                    else:
                        remove_element(paragraph)

            i = block_end

        self._rebuild_anchors()

    def _strip_cell_numPr(self, cell_elem):
        """Remove <w:numPr> from all paragraphs in a table cell to prevent
        multi-level list numbering from leaking into table content."""
        for p in cell_elem.findall(qn("w:p")):
            pPr = p.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    pPr.remove(numPr)

    def _replace_table_element(self, tbl_elem, visual: dict):
        """Replace data in a question data table element."""
        rows = tbl_elem.findall(qn("w:tr"))
        table_data = visual.get("table_data", {})
        options = table_data.get("options", [])

        if len(rows) < 3:
            return

        # Row 0: Question + option codes
        # Row 1: Sample counts
        # Row 2: Percentages

        for row_idx in range(len(rows)):
            cells = rows[row_idx].findall(qn("w:tc"))
            if row_idx == 0:
                # Header row: question text in first cell, option labels in others
                if len(cells) >= 2:
                    # First cell = question text
                    self._strip_cell_numPr(cells[0])
                    for p in cells[0].findall(qn("w:p")):
                        if _get_paragraph_text(p).strip():
                            _set_paragraph_text(p, table_data.get("question", ""))
                    # Option cells
                    for ci, opt in enumerate(options):
                        cell_i = ci + 2
                        if cell_i < len(cells):
                            self._strip_cell_numPr(cells[cell_i])
                            for p in cells[cell_i].findall(qn("w:p")):
                                for r in p.findall(qn("w:r")):
                                    for t in r.findall(qn("w:t")):
                                        t.text = f"{opt.get('code', '')}.{opt.get('text', '')}"
                                        t.set(qn("xml:space"), "preserve")
            elif row_idx == 1:
                # Count row
                if len(cells) >= 2:
                    for ci, opt in enumerate(options):
                        cell_i = ci + 2
                        if cell_i < len(cells):
                            self._strip_cell_numPr(cells[cell_i])
                            for p in cells[cell_i].findall(qn("w:p")):
                                for r in p.findall(qn("w:r")):
                                    for t in r.findall(qn("w:t")):
                                        t.text = str(opt.get("count", ""))
                                        t.set(qn("xml:space"), "preserve")
            elif row_idx == 2:
                # Percentage row
                for ci, opt in enumerate(options):
                    cell_i = ci + 2
                    if cell_i < len(cells):
                        self._strip_cell_numPr(cells[cell_i])
                        for p in cells[cell_i].findall(qn("w:p")):
                            for r in p.findall(qn("w:r")):
                                for t in r.findall(qn("w:t")):
                                    t.text = str(opt.get("pct", ""))
                                    t.set(qn("xml:space"), "preserve")

    def replace_summary(self):
        """Replace 调研结果 summary sections."""
        summary = self.payload.get("summary", {})

        self._ensure_summary_headings()
        self._rebuild_anchors()

        # 5.1 问卷重点问题分析
        key_issue_items = summary.get("key_issue_items", [])
        if key_issue_items:
            self._replace_key_issue_section(key_issue_items)

        # 5.2 调研结果分析
        overall = summary.get("overall_analysis", [])
        if overall:
            self._replace_section_body_by_heading(
                "5.2调研结果分析",
                "5.3建议",
                overall,
            )

        # 5.3 建议
        recommendations = summary.get("recommendations", [])
        if recommendations:
            self._replace_section_body_by_heading(
                "5.3建议",
                "附件1",
                recommendations,
            )

    def _ensure_summary_headings(self):
        """Ensure 调研结果 / 5.1 / 5.2 / 5.3 headings exist before replacement."""
        if self._find_exact_paragraph_index("调研结果") is not None and self._find_exact_paragraph_index("5.1问卷重点问题分析") is not None:
            return

        anchor_idx = self._find_exact_paragraph_index("5.2调研结果分析")
        if anchor_idx is None:
            return
        ref_child = self._get_paragraph_at(anchor_idx)

        for text in ["调研结果", "5.1问卷重点问题分析"]:
            if self._find_exact_paragraph_index(text) is not None:
                continue
            new_p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            new_p.append(r)
            ref_child.addprevious(new_p)
            _set_paragraph_style_props(new_p, "宋体", 16, True, "left")

    def _make_chart_png(self, chart: dict, output_path: Path, role: str = "overview") -> Optional[Path]:
        categories = [str(value) for value in chart.get("categories", []) if str(value)]
        values = [float(value) for value in chart.get("values", [])]
        if not categories or not values:
            return None
        item_count = min(len(categories), len(values))
        categories = categories[:item_count]
        values = values[:item_count]

        import os
        cache_root = Path(tempfile.gettempdir()) / "jlk_mpl_cache"
        cache_root.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("MPLCONFIGDIR", str(cache_root / "mplconfig"))
        os.environ.setdefault("XDG_CACHE_HOME", str(cache_root / "xdg"))
        os.environ.setdefault("FC_CACHEDIR", str(cache_root / "fontconfig"))
        Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
        Path(os.environ["XDG_CACHE_HOME"]).mkdir(parents=True, exist_ok=True)
        Path(os.environ["FC_CACHEDIR"]).mkdir(parents=True, exist_ok=True)
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "PingFang SC", "Heiti TC", "SimHei", "DejaVu Sans"]
        plt.rcParams["axes.unicode_minus"] = False

        title = chart.get("title", "")
        chart_type = chart.get("chart_type", "bar")
        blue = "#5aa9e6"
        pie_colors = ["#5b57f2", "#58d99a", "#ffda4d", "#ff9a40", "#5aa9e6", "#66d0a4", "#617594", "#a9d18e"]
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if chart_type == "pie":
            fig, ax = plt.subplots(figsize=(6.90, 4.51), dpi=160)
            colors = pie_colors[:len(categories)]
            if role == "overview":
                wedges, _ = ax.pie(
                    values,
                    startangle=90,
                    counterclock=False,
                    colors=colors,
                    radius=0.86,
                    wedgeprops={"linewidth": 1.0, "edgecolor": "white"},
                )
                total = sum(values)
                for wedge, value, color in zip(wedges, values, colors):
                    if not total:
                        continue
                    percent = value / total * 100
                    angle = math.radians((wedge.theta1 + wedge.theta2) / 2)
                    x, y = math.cos(angle), math.sin(angle)
                    start = (0.86 * x, 0.86 * y)
                    mid = (1.02 * x, 1.02 * y)
                    end = (1.20 * (1 if x >= 0 else -1), 1.02 * y)
                    ax.plot([start[0], mid[0], end[0]], [start[1], mid[1], end[1]], color=color, linewidth=1.4)
                    ax.text(
                        end[0] + (0.04 if x >= 0 else -0.04),
                        end[1],
                        f"{percent:.0f}%",
                        ha="left" if x >= 0 else "right",
                        va="center",
                        fontsize=12,
                        color="#777777",
                    )
                handles = [
                    plt.Line2D([0], [0], marker="o", linestyle="", markersize=8, markerfacecolor=color, markeredgecolor=color)
                    for color in colors
                ]
                ax.legend(
                    handles,
                    categories,
                    loc="lower center",
                    bbox_to_anchor=(0.5, -0.12),
                    ncol=4,
                    frameon=False,
                    fontsize=10,
                    handlelength=0.8,
                    columnspacing=1.1,
                    handletextpad=0.4,
                )
                ax.set_xlim(-1.45, 1.45)
                ax.set_ylim(-1.28, 1.25)
            else:
                depth = 0.08
                dark_colors = [_darken_hex(color, 0.72) for color in colors]
                for layer in range(8, 0, -1):
                    ax.pie(
                        values,
                        startangle=90,
                        colors=dark_colors,
                        radius=0.82,
                        center=(0, 0.04 - depth * layer / 8),
                        wedgeprops={"linewidth": 0.5, "edgecolor": "white"},
                    )
                ax.pie(
                    values,
                    labels=categories,
                    autopct=lambda pct: f"{pct:.1f}%" if pct >= 3 else "",
                    startangle=90,
                    colors=colors,
                    radius=0.82,
                    center=(0, 0.04),
                    pctdistance=0.62,
                    labeldistance=1.18,
                    textprops={"fontsize": 8.5, "color": "#111111"},
                    wedgeprops={"linewidth": 0.8, "edgecolor": "white"},
                )
                if title:
                    ax.set_title(title, fontsize=12, pad=8)
            ax.axis("equal")
            fig.subplots_adjust(left=0.04, right=0.96, top=0.90, bottom=0.18)
        else:
            fig_height = max(4.05, 0.48 * len(categories) + 0.95)
            fig, ax = plt.subplots(figsize=(6.93, fig_height), dpi=160)
            y_pos = list(range(len(categories)))
            ax.barh(y_pos, values, color=blue, edgecolor=blue, height=0.56, label="值")
            ax.set_yticks(y_pos)
            ax.set_yticklabels(categories, fontsize=11, color="#666666")
            ax.invert_yaxis()
            ax.grid(axis="x", color="#e6e6e6", linewidth=1.0)
            ax.set_axisbelow(True)
            for spine in ["top", "right"]:
                ax.spines[spine].set_visible(False)
            ax.spines["left"].set_color("#bfbfbf")
            ax.spines["bottom"].set_color("#bfbfbf")
            ax.tick_params(axis="y", length=0)
            ax.tick_params(axis="x", colors="#777777", labelsize=10)
            max_value = max(float(v) for v in values) if values else 0
            ax.set_xlim(0, max_value * 1.08 if max_value else 1)
            for y, value in enumerate(values):
                ax.text(float(value) + (max_value * 0.012 if max_value else 0.02), y, f"{value:g}", va="center", fontsize=11, color="#666666")
            ax.legend(loc="lower center", bbox_to_anchor=(0.5, -0.22), frameon=False, fontsize=10, handlelength=0.8)
            if title and role != "overview":
                ax.set_title(title, fontsize=12, pad=10)
            fig.subplots_adjust(left=0.28, right=0.96, top=0.96, bottom=0.20)

        fig.savefig(output_path, facecolor="white")
        plt.close(fig)
        return output_path

    def _replace_paragraph_with_image(self, paragraph_element, image_path: Path, width_inches: float = 5.75) -> None:
        for run in list(paragraph_element.findall(qn("w:r"))):
            paragraph_element.remove(run)
        paragraph = Paragraph(paragraph_element, self.doc._body)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))

    def replace_native_charts_with_pngs(self) -> None:
        result_analysis = self.payload.get("result_analysis", {})
        overview_charts = result_analysis.get("overview_charts", [])
        if not overview_charts:
            return

        chart_dir = Path(tempfile.mkdtemp(prefix="jlk_charts_"))
        overview_pngs = []
        for idx, chart in enumerate(overview_charts[:2]):
            png = self._make_chart_png(chart, chart_dir / f"overview_{idx + 1}.png", role="overview")
            if png:
                overview_pngs.append(png)

        drawing_paragraphs = []
        found_intro = False
        for child in self.body:
            if child.tag.split("}")[-1] == "sectPr":
                break
            if child.tag.split("}")[-1] != "p":
                continue
            text = _get_paragraph_text(child).strip()
            if "数据可视化" in text:
                found_intro = True
                continue
            if found_intro and re.match(r"^4\.\d+", text):
                break
            if found_intro and _paragraph_has_drawing(child):
                drawing_paragraphs.append(child)

        native_pie_template = None
        for paragraph in drawing_paragraphs:
            if not _paragraph_has_chart(paragraph):
                continue
            native_pie_template = _clone_first_chart_paragraph(paragraph, "rIdJlkKeyIssueChart1")
            if native_pie_template is not None:
                break

        for idx, png in enumerate(overview_pngs):
            if idx < len(drawing_paragraphs):
                self._replace_paragraph_with_image(drawing_paragraphs[idx], png)

        for extra in drawing_paragraphs[len(overview_pngs):]:
            if extra.getparent() is not None:
                extra.getparent().remove(extra)

        key_issue_items = self.payload.get("summary", {}).get("key_issue_items", [])
        if not key_issue_items:
            self._rebuild_anchors()
            return

        self._rebuild_anchors()
        start_idx = self._find_exact_paragraph_index("5.1问卷重点问题分析")
        end_idx = self._find_exact_paragraph_index("5.2调研结果分析", start=(start_idx + 1) if start_idx is not None else 0)
        if start_idx is None or end_idx is None or end_idx <= start_idx:
            self._rebuild_anchors()
            return

        if native_pie_template is not None:
            insert_before = self._get_paragraph_at(start_idx + 1) if start_idx + 1 < end_idx else self._get_paragraph_at(end_idx)
            for idx in range(min(2, len(key_issue_items))):
                chart_paragraph = copy.deepcopy(native_pie_template)
                for chart in chart_paragraph.findall(".//c:chart", NS):
                    chart.set(qn("r:id"), f"rIdJlkKeyIssueChart{idx + 1}")
                insert_before.addprevious(chart_paragraph)

        self._rebuild_anchors()

    def _replace_key_issue_section(self, key_issue_items: list[dict]):
        start_idx = self._find_exact_paragraph_index("5.1问卷重点问题分析")
        end_idx = self._find_exact_paragraph_index("5.2调研结果分析", start=(start_idx + 1) if start_idx is not None else 0)
        if start_idx is None or end_idx is None or end_idx <= start_idx:
            return

        text_paragraphs = []
        for i in range(start_idx + 1, end_idx):
            info = self.anchors.get(i)
            if not info or info["type"] != "p":
                continue
            paragraph = info["elem"]
            if _paragraph_has_drawing(paragraph):
                continue
            text_paragraphs.append(paragraph)

        replacement_paragraphs = [item.get("paragraph", "") for item in key_issue_items]

        insert_before = self._get_paragraph_at(end_idx)
        for pi, text in enumerate(replacement_paragraphs):
            if pi < len(text_paragraphs):
                _set_paragraph_text(text_paragraphs[pi], text)
                _set_paragraph_style_props(
                    text_paragraphs[pi],
                    "宋体",
                    12,
                    False,
                    "both",
                    body_layout=True,
                    first_line_chars=200,
                )
            else:
                new_p = OxmlElement("w:p")
                new_r = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.set(qn("xml:space"), "preserve")
                new_t.text = text
                new_r.append(new_t)
                new_p.append(new_r)
                insert_before.addprevious(new_p)
                _set_paragraph_style_props(
                    new_p,
                    "宋体",
                    12,
                    False,
                    "both",
                    body_layout=True,
                    first_line_chars=200,
                )

        for extra in text_paragraphs[len(replacement_paragraphs):]:
            extra.getparent().remove(extra)

        self._rebuild_anchors()

    def _replace_section_body_by_heading(self, start_heading: str, end_heading: str, new_paragraphs: list[str]):
        """Replace paragraph content between two exact heading paragraphs."""
        start_idx = self._find_exact_paragraph_index(start_heading)
        end_idx = self._find_exact_paragraph_index(end_heading, start=(start_idx + 1) if start_idx is not None else 0)
        if start_idx is None or end_idx is None or end_idx <= start_idx:
            return

        body_paragraphs = []
        for i in range(start_idx + 1, end_idx):
            info = self.anchors.get(i)
            if not info or info["type"] != "p":
                continue
            body_paragraphs.append(info["elem"])

        insert_before = self._get_paragraph_at(end_idx)
        for pi, text in enumerate(new_paragraphs):
            if pi < len(body_paragraphs):
                _set_paragraph_text(body_paragraphs[pi], text)
                _set_paragraph_style_props(body_paragraphs[pi], "宋体", 12, False)
            else:
                new_p = OxmlElement("w:p")
                new_r = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.set(qn("xml:space"), "preserve")
                new_t.text = text
                new_r.append(new_t)
                new_p.append(new_r)
                insert_before.addprevious(new_p)
                _set_paragraph_style_props(new_p, "宋体", 12, False)

        for extra in body_paragraphs[len(new_paragraphs):]:
            extra.getparent().remove(extra)

        self._rebuild_anchors()

    def _replace_paragraphs_after(self, start_idx: int, new_paragraphs: list[str], stop_patterns: list[str] = None):
        """Replace paragraphs after a heading with new text."""
        stop_patterns = stop_patterns or []
        target_paragraphs = []

        for i in range(start_idx + 1, start_idx + 40):
            if i not in self.anchors:
                break
            info = self.anchors[i]
            if info["type"] != "p":
                continue
            text = info["text"]
            if any(sp in text for sp in stop_patterns):
                break
            if text.strip():
                target_paragraphs.append(info["elem"])

        for pi, paragraph in enumerate(target_paragraphs):
            if pi < len(new_paragraphs):
                _set_paragraph_text(paragraph, new_paragraphs[pi])
                _set_paragraph_style_props(paragraph, "宋体", 12, False)
            else:
                paragraph.getparent().remove(paragraph)

    def replace_attachments(self):
        """Replace attachment 1 with actual question list."""
        attachments = self.payload.get("attachments", {})
        questions = attachments.get("attachment1_questions", [])
        if not questions:
            return

        idx = self._find_anchor_index("附件1")
        if idx is None:
            return

        att_name = attachments.get("attachment1_name", "")
        if att_name:
            p = self._get_paragraph_at(idx)
            _set_paragraph_text(p, f"附件1：{att_name}")
            _set_paragraph_style_props(p, "宋体", 16, True, "left")

        attachment2_idx = self._find_anchor_index("附件2", start=idx + 1)
        if attachment2_idx is None:
            return

        desired: list[str] = []
        for display_index, question in enumerate(questions, start=1):
            question_text = re.sub(r"^\s*\d+\s*[\.．、]\s*", "", str(question.get("question", "")).strip())
            desired.append(f"（{display_index}） {question_text}")
            for option in question.get("options", []):
                desired.append(f"{option.get('code', '')}. {option.get('text', '')}")

        target_paragraphs = []
        for i in range(idx + 1, attachment2_idx):
            info = self.anchors.get(i)
            if not info or info["type"] != "p":
                continue
            paragraph = info["elem"]
            if _paragraph_has_drawing(paragraph):
                continue
            if _get_paragraph_text(paragraph).strip():
                target_paragraphs.append(paragraph)

        if not target_paragraphs:
            return

        insert_before = self.anchors[attachment2_idx]["elem"]
        template_paragraph = target_paragraphs[-1]
        while len(target_paragraphs) < len(desired):
            cloned = copy.deepcopy(template_paragraph)
            _strip_numPr(cloned)
            insert_before.addprevious(cloned)
            target_paragraphs.append(cloned)

        for index, paragraph in enumerate(target_paragraphs):
            _strip_numPr(paragraph)
            if index < len(desired):
                _set_paragraph_text(paragraph, desired[index])
                _set_paragraph_style_props(paragraph, "宋体", 12, False)
            else:
                parent = paragraph.getparent()
                if parent is not None:
                    parent.remove(paragraph)

        self._rebuild_anchors()

    def replace_disclaimer(self):
        """Replace disclaimer section."""
        disclaimer = self.payload.get("disclaimer", {})
        items = disclaimer.get("items", [])
        unit = self.payload.get("service", {}).get("unit", disclaimer.get("unit", ""))
        date_str = self.payload.get("service", {}).get("date", disclaimer.get("date", ""))

        # Find "免责申明" heading
        idx = self._find_anchor_index("免责申明")
        if idx is None:
            return
        heading = self._get_paragraph_at(idx)
        _set_paragraph_style_props(heading, "宋体", 16, True, "center")

        # Replace disclaimer items and right-aligned signature lines.
        item_count = 0
        unit_replaced = False
        date_replaced = False
        for i in range(idx + 1, idx + 40):
            if i not in self.anchors:
                break
            info = self.anchors[i]
            if info["type"] != "p":
                continue
            text = info["text"].strip()
            if "服务提供单位" in text:
                p = self._get_paragraph_at(i)
                _set_paragraph_text(p, f"服务提供单位:{unit}")
                _set_paragraph_style_props(p, "宋体", 12, False, "right", body_layout=True, first_line_chars=0)
                unit_replaced = True
                continue
            if re.match(r"20\d{2}年", text):
                p = self._get_paragraph_at(i)
                _set_paragraph_text(p, date_str)
                _set_paragraph_style_props(p, "宋体", 12, False, "right", body_layout=True, first_line_chars=0)
                date_replaced = True
                continue
            if text.startswith("（") and item_count < len(items):
                p = self._get_paragraph_at(i)
                _set_paragraph_text(p, items[item_count])
                _set_paragraph_style_props(p, "宋体", 12, False, "both", body_layout=True, first_line_chars=0)
                item_count += 1

    def restyle_key_issue_titles(self):
        """Force 5.1 custom sub-headings back to title layout after body formatting."""
        for item in self.payload.get("summary", {}).get("key_issue_items", []):
            heading_text = item.get("heading")
            if not heading_text:
                continue
            idx = self._find_exact_paragraph_index(heading_text)
            if idx is None:
                continue
            paragraph = self._get_paragraph_at(idx)
            _set_paragraph_style_props(paragraph, "宋体", 12, True, "left", body_layout=False)

    # ─── Paragraph formatting ───────────────────────────────────────────
    
    def _apply_uniform_paragraph_formatting(self):
        """Apply uniform paragraph formatting to all body text paragraphs.
        
        Settings (Chinese document standard):
        - Alignment: justified (两端对齐)
        - Outline level: body text (正文文本)
        - Indent: first-line 2 chars, 0 before/after
        - Spacing: line 2.5x multiple, 0 before/after
        - Snap to grid: on
        
        Skips: heading paragraphs (pStyle 2/3/4), TOC, special-layout
        paragraphs (center/right aligned — title page, service info).
        """
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        WNS = f"{{{W}}}"
        
        skipped_headings = set()
        skipped_special = []
        formatted_count = 0
        
        for child in self.body:
            tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                break
            if tag != "p":
                continue
            
            pPr = child.find(f"{WNS}pPr")
            if pPr is None:
                # No pPr at all — create one (likely an injected paragraph)
                from lxml import etree
                pPr = etree.SubElement(child, f"{WNS}pPr")
            
            # ── Skip headings (pStyle 2/3/4) ──
            pStyle = pPr.find(f"{WNS}pStyle")
            if pStyle is not None:
                val = pStyle.get(f"{WNS}val", "")
                if val in ("2", "3", "4"):
                    skipped_headings.add(val)
                    continue
            
            text = _get_paragraph_text(child).strip()
            if (
                not text
                or text == "目录"
                or "更新域" in text
                or text == "问卷调研服务结算"
                or text == "前言"
                or text.startswith("服务单位：")
                or text.startswith("日期：")
                or text == self.payload.get("report_title", "")
            ):
                skipped_special.append(text[:20] or "<empty>")
                continue

            _apply_body_paragraph_layout(child, align="both")
            formatted_count += 1
        
        # Summary
        parts = []
        if formatted_count:
            parts.append(f"{formatted_count} body paragraphs")
        if skipped_headings:
            parts.append(f"skipped heading levels {sorted(skipped_headings)}")
        if skipped_special:
            parts.append(f"skipped special paragraphs {len(skipped_special)}")
        if parts:
            print(f"[formatting] Applied uniform format: {', '.join(parts)}")

    def _apply_global_font_name(self):
        """Normalize visible run font names without changing size or weight."""
        for run in self.body.iter(qn("w:r")):
            if not "".join(t.text or "" for t in run.findall(qn("w:t"))).strip():
                continue
            rPr = run.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "宋体")
            rFonts.set(qn("w:hAnsi"), "宋体")
            rFonts.set(qn("w:eastAsia"), "宋体")

    def _apply_red_font_replacements(self):
        """Apply red-font variable replacements across ALL text in the document.
        
        This is a final cleanup pass that ensures any hardcoded template text
        (product names, regions, sample sizes) is replaced even in paragraphs
        generated by AI that still contain old template values.
        """
        rf = self.payload.get("red_font_replacements", {})
        if not rf:
            return
        
        # Build replacement map: old_text -> new_text
        # Keys are the template's original values, values are the new ones
        replacements = {}
        product = rf.get("drug_name", "")
        region = rf.get("region", "")
        sample_size = rf.get("sample_size", "")
        
        if product:
            replacements["厄贝沙坦氢氯噻嗪片"] = product
        if region:
            replacements["广东省"] = region
        if sample_size:
            replacements["1642份"] = f"{sample_size}份"
            replacements["1642名"] = f"{sample_size}名"
            replacements["1642"] = sample_size  # standalone number
        
        if not replacements:
            return
        
        # Apply to all paragraphs in body
        for child in self.body:
            tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if tag == "sectPr":
                break
            if tag == "p":
                self._safe_replace_multi(child, replacements)
            elif tag == "tbl":
                for tc in child.iter(qn("w:tc")):
                    for p in tc.findall(qn("w:p")):
                        self._safe_replace_multi(p, replacements)
        
        # Apply to headers and footers
        for section in self.doc.sections:
            for container in [section.header, section.footer,
                              section.first_page_header, section.first_page_footer]:
                try:
                    for p in container.paragraphs:
                        self._safe_replace_multi(p._element, replacements)
                except:
                    pass
    
    def _safe_replace_multi(self, p_element, replacements: dict):
        """Replace multiple strings in a paragraph, preserving drawings."""
        has_drawing = any(
            r.find(qn("w:drawing")) is not None
            for r in p_element.findall(qn("w:r"))
        )
        
        if has_drawing:
            # Only replace in non-drawing runs
            for r in p_element.findall(qn("w:r")):
                if r.find(qn("w:drawing")) is not None:
                    continue
                for t in r.findall(qn("w:t")):
                    if t.text:
                        for old, new in replacements.items():
                            if old in t.text:
                                t.text = t.text.replace(old, new)
                                t.set(qn("xml:space"), "preserve")
        else:
            # Full paragraph replacement
            full = _get_paragraph_text(p_element)
            new_full = full
            for old, new in replacements.items():
                new_full = new_full.replace(old, new)
            if new_full != full:
                _set_paragraph_text(p_element, new_full)

    # ─── Chart update via ZIP manipulation ─────────────────────────────

    def update_charts(self, output_docx: Path):
        """Update Office-native chart data in the output docx.

        This modifies chart XML directly inside the ZIP to preserve styling.
        Called AFTER the doc is initially saved.
        """
        # Determine which chart to update from payload
        result_analysis = self.payload.get("result_analysis", {})
        sections = result_analysis.get("sections", [])

        # Chart 1/2 are the result-analysis overview charts in the template.

        if not sections:
            return

        q_lookup = {}
        for sec in sections:
            for vg in sec.get("visual_groups", []):
                table_data = vg.get("table_data", {})
                question_ref = vg.get("question_ref")
                if question_ref and table_data:
                    q_lookup[question_ref] = table_data

        overview_charts = result_analysis.get("overview_charts", [])
        overview_bindings = [("word/charts/chart1.xml", 0), ("word/charts/chart2.xml", 1)]
        for chart_path, chart_index in overview_bindings:
            if chart_index >= len(overview_charts):
                continue
            overview = overview_charts[chart_index]
            categories = overview.get("categories", [])
            values = overview.get("values", [])
            if categories and values:
                self._update_chart_in_zip(
                    output_docx,
                    chart_path,
                    categories,
                    values,
                    overview.get("title") or "问卷结果分析维度占比图",
                )

    def _update_chart_in_zip(self, docx_path: Path, chart_path_in_zip: str, categories: list, values: list, title: str):
        """Extract, modify, and replace chart XML inside the docx ZIP."""
        tmp_path = docx_path.with_suffix(".tmp.docx")

        with zipfile.ZipFile(docx_path, "r") as zin:
            names = set(zin.namelist())
            source_chart = chart_path_in_zip if chart_path_in_zip in names else "word/charts/chart1.xml"
            chart_data = zin.read(source_chart)
            with tempfile.NamedTemporaryFile(suffix=".xml", delete=False) as tf:
                tf.write(chart_data)
                tf_path = Path(tf.name)
            try:
                _update_chart_data(tf_path, categories, values, title)
                updated_chart_data = tf_path.read_bytes()
            finally:
                tf_path.unlink()

            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                wrote_chart = False
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == chart_path_in_zip:
                        data = updated_chart_data
                        wrote_chart = True
                    zout.writestr(item, data)
                if not wrote_chart:
                    zout.writestr(chart_path_in_zip, updated_chart_data)

        docx_path.unlink()
        tmp_path.rename(docx_path)

    def _key_issue_chart_values(self, values: list) -> list[float]:
        numeric = [float(value) for value in values]
        if numeric and sum(numeric) > 1.5:
            return [value / 100 for value in numeric]
        return numeric

    def _build_native_chart_xml(self, source_xml: bytes, title: str, categories: list[str], values: list[float]) -> bytes:
        with tempfile.NamedTemporaryFile(suffix=".xml", delete=False) as tf:
            tf.write(source_xml)
            temp_path = Path(tf.name)
        try:
            _update_chart_data(temp_path, categories, values, title)
            return temp_path.read_bytes()
        finally:
            temp_path.unlink(missing_ok=True)

    def _build_native_chart_workbook(self, categories: list[str], values: list[float]) -> bytes:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Sheet1"
        worksheet["A1"] = " "
        worksheet["B1"] = "占比"
        for index, (category, value) in enumerate(zip(categories, values), start=2):
            worksheet.cell(row=index, column=1).value = category
            value_cell = worksheet.cell(row=index, column=2)
            value_cell.value = value
            value_cell.number_format = "0.00%"
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()

    def _add_content_type_override(self, content_types_xml: bytes, part_name: str, content_type: str) -> bytes:
        ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
        root = ET.fromstring(content_types_xml)
        for override in root.findall("ct:Override", ns):
            if override.attrib.get("PartName") == part_name:
                return content_types_xml
        override = ET.Element(f"{{{ns['ct']}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", content_type)
        root.append(override)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _add_document_chart_relationships(self, rels_xml: bytes) -> bytes:
        rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        chart_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
        root = ET.fromstring(rels_xml)
        existing = {rel.attrib.get("Id") for rel in root}
        for index in (1, 2):
            rel_id = f"rIdJlkKeyIssueChart{index}"
            if rel_id in existing:
                continue
            rel = ET.Element(f"{{{rel_ns}}}Relationship")
            rel.set("Id", rel_id)
            rel.set("Type", chart_type)
            rel.set("Target", f"charts/chart{index + 2}.xml")
            root.append(rel)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _build_chart_relationships(self, source_rels: bytes, workbook_name: str) -> bytes:
        rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        package_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"
        root = ET.fromstring(source_rels)
        for rel in root:
            if rel.attrib.get("Type") == package_type:
                rel.set("Target", f"../embeddings/{workbook_name}")
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _write_native_key_issue_chart_parts(self, docx_path: Path):
        key_issue_items = self.payload.get("summary", {}).get("key_issue_items", [])[:2]
        if not key_issue_items:
            return

        tmp_path = docx_path.with_suffix(".keycharts.tmp.docx")
        chart_content_type = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
        with zipfile.ZipFile(docx_path, "r") as zin:
            source_chart = zin.read("word/charts/chart1.xml")
            source_chart_rels = zin.read("word/charts/_rels/chart1.xml.rels")
            generated: dict[str, bytes] = {
                "word/_rels/document.xml.rels": self._add_document_chart_relationships(
                    zin.read("word/_rels/document.xml.rels")
                ),
                "[Content_Types].xml": zin.read("[Content_Types].xml"),
            }

            for index, item in enumerate(key_issue_items, start=1):
                chart_name = f"word/charts/chart{index + 2}.xml"
                workbook_name = f"Workbook{index + 3}.xlsx"
                workbook_path = f"word/embeddings/{workbook_name}"
                chart_rels_name = f"word/charts/_rels/chart{index + 2}.xml.rels"
                categories = [str(value) for value in item.get("categories", [])]
                values = self._key_issue_chart_values(item.get("values", []))
                generated[chart_name] = self._build_native_chart_xml(
                    source_chart,
                    item.get("chart_title", ""),
                    categories,
                    values,
                )
                generated[workbook_path] = self._build_native_chart_workbook(categories, values)
                generated[chart_rels_name] = self._build_chart_relationships(source_chart_rels, workbook_name)
                generated["[Content_Types].xml"] = self._add_content_type_override(
                    generated["[Content_Types].xml"],
                    f"/{chart_name}",
                    chart_content_type,
                )

            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                replaced = set()
                for item in zin.infolist():
                    if item.filename in generated:
                        zout.writestr(item, generated[item.filename])
                        replaced.add(item.filename)
                    else:
                        zout.writestr(item, zin.read(item.filename))
                for filename, data in generated.items():
                    if filename not in replaced:
                        zout.writestr(filename, data)

        docx_path.unlink()
        tmp_path.rename(docx_path)

    def _enable_update_fields_on_open(self, docx_path: Path):
        """Ensure Word refreshes fields such as TOC on document open."""
        tmp_path = docx_path.with_suffix(".settings.tmp.docx")
        settings_path = "word/settings.xml"

        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == settings_path:
                        root = ET.fromstring(data)
                        update = root.find(qn("w:updateFields"))
                        if update is None:
                            update = ET.Element(qn("w:updateFields"))
                            root.append(update)
                        update.set(qn("w:val"), "true")
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                    zout.writestr(item, data)

        docx_path.unlink()
        tmp_path.rename(docx_path)

    def _normalize_font_references(self, docx_path: Path):
        """Normalize document/style/theme font declarations for cross-platform Word rendering."""
        tmp_path = docx_path.with_suffix(".fonts.tmp.docx")
        font_names = [
            "宋体",
            "黑体",
            "微软雅黑",
            "汉仪中宋简",
            "Times New Roman",
            "Arial",
            "Calibri",
            "SimSun",
            "SimHei",
        ]

        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.endswith(".xml"):
                        text = data.decode("utf-8", errors="replace")
                        if item.filename.startswith("word/"):
                            for name in font_names:
                                text = text.replace(name, DEFAULT_FONT)
                            text = re.sub(
                                r'w:(ascii|hAnsi|eastAsia|cs)(Theme)?="[^"]*"',
                                lambda match: f'w:{match.group(1)}="{DEFAULT_FONT}"',
                                text,
                            )
                            if item.filename == "word/theme/theme1.xml":
                                text = re.sub(r'typeface="[^"]*"', f'typeface="{DEFAULT_FONT}"', text)
                        data = text.encode("utf-8")
                    zout.writestr(item, data)

        docx_path.unlink()
        tmp_path.rename(docx_path)

    # ─── Main render ────────────────────────────────────────────────────

    def render(self, output_path: Path) -> Path:
        """Execute all replacements and save."""
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Phase 1: Simple text replacements (in order of template structure)
        self.replace_header()
        self.replace_settlement()
        self.replace_service_unit()
        self.replace_toc()           # structural: rebuilds anchors inside
        self.replace_metadata()
        self.replace_preface()       # structural: inserts paragraphs
        self._rebuild_anchors()
        self.replace_report_title()
        self.replace_project_background()
        self._rebuild_anchors()
        self.replace_project_execution()
        self._rebuild_anchors()
        self.replace_questionnaire_note()
        self.replace_result_analysis_intro()
        self.replace_analysis_sections()  # uses children list, no index dependency
        self._rebuild_anchors()
        self.replace_summary()
        self._rebuild_anchors()
        self.replace_native_charts_with_pngs()
        self._rebuild_anchors()
        self.replace_attachments()
        
        # Phase 1.4: Uniform paragraph formatting (all body text: 两端对齐,
        # 首行缩进2字符, 2.5倍行距 — before red-font pass so formatted text
        # still gets variable replacement)
        self._apply_uniform_paragraph_formatting()
        
        # Phase 1.5: Global red-font replacement (catches any AI-generated text
        # that still contains old template values like "广东省" or "1642份")
        self._apply_red_font_replacements()

        # Phase 1.6: Re-apply local layouts that must override the generic
        # body-format pass.
        self._rebuild_anchors()
        self.restyle_key_issue_titles()
        self._rebuild_anchors()
        self.replace_disclaimer()
        self._apply_global_font_name()

        # Save
        self.doc.save(str(output_path))

        # Phase 2: Global XML-level text cleanup (catch anything missed)
        self._global_replace_in_xml(output_path)
        self._write_native_key_issue_chart_parts(output_path)

        # Phase 3: Enable field updates on open and normalize final XML.
        self._enable_update_fields_on_open(output_path)
        self._normalize_font_references(output_path)

        return output_path

    def _global_replace_in_xml(self, docx_path: Path):
        """Do a raw XML-level text replacement as a final cleanup pass.
        
        This catches text in containers not reached by python-docx traversal
        (e.g., SDTs, content controls, embedded text boxes).
        """
        region = self.meta.get("region", "")
        product = self.meta.get("product", "")
        survey_period = self.meta.get("survey_period", "")
        if not region and not product and not survey_period:
            return
        
        # Map old → new from the template's original data
        replacements = {}
        if region:
            replacements["广东省"] = region
        if product and product != "厄贝沙坦氢氯噻嗪片":
            replacements["厄贝沙坦氢氯噻嗪片"] = product
        sample_size = self.meta.get("sample_size") or self.meta.get("valid_count")
        if sample_size:
            replacements["1642"] = str(sample_size)
        
        tmp_path = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                        text = data.decode("utf-8", errors="replace")
                        for old, new in replacements.items():
                            text = text.replace(old, new)
                        if survey_period and item.filename == "word/document.xml":
                            text = re.sub(r"调研时间：[^<]+", f"调研时间：{survey_period}", text)
                        data = text.encode("utf-8")
                    zout.writestr(item, data)
        
        docx_path.unlink()
        tmp_path.rename(docx_path)


# ─── CLI ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Render a patient report from template using object-level replacement."
    )
    parser.add_argument("payload_json", help="Path to report_payload.json")
    parser.add_argument("-o", "--output", required=True, help="Output docx path")
    args = parser.parse_args()

    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))

    # Template path from payload or default
    template_path = payload.get("meta", {}).get("template_doc")
    if not template_path:
        # Default to the template bundled with the skill
        skill_root = Path(__file__).resolve().parent.parent
        template_path = skill_root / "templates" / "efficacy-report-template.docx"
    else:
        template_path = Path(template_path)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    output = Path(args.output)

    renderer = TemplateRenderer(template_path, payload)
    result = renderer.render(output)
    print(result)


if __name__ == "__main__":
    import re
    main()
